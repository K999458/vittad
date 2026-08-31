#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_txxb_docx.py —— 《图学学报》(TXXB) 投稿初稿构建脚本（TXXB_FORK_MARK）

由 build_oep_docx.py v12 fork。转刊改造（对照 txxb_tpl_submit 模板、投稿指南
与两篇已刊文章 DOI 10.11996/JG.j.2095-302X.2026010029/2026020380 实测版式）：
  页面 A4 边距 T2.8/B1.7/L2.0/R2.0，正文行距固定 15.7 磅；引言不设标题不编号
  （已刊惯例），正文编号 1 起；头区两行标签（中图分类号 TP 391/DOI/文献标识码/
  文章编号）；作者 4 号仿宋、单位小 5 仿宋、摘要关键词小 5 楷体两边缩 2 字符；
  英文题 4 号 TNR 加粗 sentence case、英文单位 6 号 TNR 加 ", China"、
  Abstract≥2 500 字符（被动语态，无 this paper）；标题一级 4 号黑体/二级 5 号
  黑体/三级 5 号楷体；图题图下黑体中文行+TNR 加粗英文行（Fig. N 带空格）、
  表题表上同式；参考文献 顺序编码制 6 号（8 磅）、标题「参考文献 (References)」；
  首页脚注区收稿日期/基金项目/作者简介双语占位；投稿双版本：原稿+盲审稿
  （删中英文作者姓名、单位及作者简介，--blind 或 build(blind=True)）。

从三类只读输入生成排版初稿 docx：
  1) ../DRAFT_v1_zh.md      —— 中文稿全文（正文表格数据在构建时直接解析，避免手工转录错误）
  2) ../references.bib      —— 55 条参考文献（构建时自动转 GB/T 7714—2015 顺序编码制）
  3) ../figures/*.png       —— 正文 8 张主图

排版方式：以官方《中文论文模板 (2024).docx》为底本复制，继承其页面设置
（A4、上下 2.15 cm/左右 1.95 cm 边距、docGrid 312）、页眉页脚（含首页
收稿日期/基金项目脚注区、奇偶页眉）与分节结构（单栏头部 + 双栏正文 +
通栏插排），清空正文后用 python-docx 按模板样式逐段重建。

输出（均写入本目录，不改动任何输入）：
  GCAS_光学精密工程_排版初稿_v8_20260829.docx
  （v2：修复 parse_bib 末字段丢失致 55 条文献年份全空的 bug；会议条目改为
  GB/T 7714—2015 会议录全称著录，新增 PROCEEDINGS_MAP；v1 文件保留不覆盖）
  （v3：对照《光学 精密工程》已刊文章体例的六项修订——①中英摘要改为四要素
  连贯段落，删去显式"目的：/方法：/结果：/结论："引导词；②图表题注改为简短
  名词短语式双语题注（图1/Fig. 1、表1/Table 1），分图与图例说明补齐中英对照两行；
  ③删除"动机。/构建。/算子。/性质。"等楷体段首标签及行内楷体强调，融入行文
  （连带修复 2.3 节星号残留）；④删除算法 1 伪代码浮动框（中文期刊无此惯例，
  流程以正文与图 1 表达），2.6 节相应改写；⑤段首粗体引导词（数据集与指标/协议/
  基线/实现/融合算子/池化规模/平滑强度/逐级归因等）全部删除并改写衔接；
  ⑥图片排版顺序复核：8 图均先文后图、随首次引用就近放置，未改动；v1/v2 保留）
  （v4：落实五路并行审查（REVIEW_v3_五路并行审查_20260828.md）中本轮可自动化的
  修复——①一式一号：8 个复合公式拆分重编号为式 (1)–(11)，全文交叉引用同步；
  ②工程实际表述落地：引言与 3.1 节补 Real3D-AD 实测采集属性（PMAX-S130 双目
  蓝光结构光扫描仪、点精度 0.011–0.015 mm、点间距 0.04–0.07 mm、360°多视角，
  数据取自其原文表 3），3.6 节 MulSen-AD 补多传感器实测属性；③全局替换：GCAS
  中文摘要首现补英文全称、fitness 正文首现补"配准适配度（fitness）"桥接、人称
  统一"本文"、千分位逗号改 GB/T 15835 式空格、直引号改全角弯引号、连接号统一
  一字线、级号统一阿拉伯数字、缩写首现补全称（AUROC/AP/kNN/normal-only/sink/
  种子）；④句级修订：五路报告 P2 全部 17 处（"第一阶段"撞名、式 (11) 回引、
  "反观各替代方案"、摘要三处、翻译腔清单等）；⑤排版：表题改小五 9pt 中文黑体、
  表内 ‡ 上标化、创新点压至约 300 字、标题编号后单空格、公式编号制表位贴栏缘、
  文章编号行改宋体、"参考文献："左对齐、EB/OL 文献补引用日期与 arXiv URL、
  Huber 补出版地；⑥新增"数据与程序可用性"段（结论后、贡献声明前，URL 占位）。
  仍留待人工/下轮：MathType 公式重排、近两年文献补引与删并、行政占位补齐、
  长句拆分与引言去重（C 级）、PDF 校样计页。v1–v3 保留）
  （v5：图 6 显示加固——作者反馈图 6（通栏 16.8 cm×约 9 cm，全文最高的通栏块）
  在其查看器中"显示不完全"。docx 内容经 XML、LibreOffice、OnlyOffice 三路核验
  本无缺损，系部分查看器（WPS/移动端/在线预览）对"连续分节中跨页高图"的兼容
  问题：图 6 落点骑跨分页边界时被拦腰截断。处置：图 6 所在单栏节由 continuous
  改为 nextPage（新页顶起排，去 titlePg 与首页页脚引用以免首页脚注区重复），
  使其在任何引擎下均整页容纳、不可能跨页；其余通栏图（图 1/图 2 块高较矮）
  机制不变。正文文字与图表数据零变化。v1–v4 保留）
  （v6：图 6 排版美化——作者反馈 v5 的 nextPage 会在前一页留整块空白。改为
  期刊通行的"通栏浮动图文框"：撤销图 6 的单双栏分节夹层，双栏正文连续排版
  （无空白），图 6+双语题注整体装入无边框单格浮动表（w:tblpPr 锚定 margin
  页顶、水平居中、tblOverlap=never 上下环绕、固定布局全文宽），浮排于所在页
  顶部并横跨双栏——与刊印论文的 [t] 通栏浮动图一致，且整体成块、仍不可能被
  分页截断。行高自适应，无文本框式裁剪风险。正文文字与图表数据零变化；
  分节 11→9，正文表格件数 7→8（新增 1 个无边框浮动容器表）。v1–v5 保留）
  （v7：空白彻底清除——作者反馈 v6 仍有空白。根因：图 1/图 2 仍沿用单双栏
  分节夹层，Word 中通栏块在页中放不下时被整体推至下页，原页留下空白带
  （空白落点随查看器分页而异）。处置：图 1/图 2 与图 6 统一改为页顶通栏
  浮动图文框（同 v6 机制），正文双栏全程连续、无分节推挤；仅表 2/3/4 区
  保留单栏夹层（三线表按行跨页续排，不产生大空白）。分节 9→5，正文表格
  件数 8→10（7 三线表 + 3 浮动容器）。新增 _qa/whitespace_scan.py 逐页空白
  检测（LibreOffice+OnlyOffice 双引擎渲染逐页量化验收）。正文文字与图表
  数据零变化。v1–v6 保留）
  （v8：栏底搁浅清除——作者指出 v7 图 6 页左栏下部仍有大块空白。根因：栏宽
  图（图 7 连题注约 10 cm 不可拆块）在本栏剩余空间放不下时整块跳往下一栏/
  页，栏底搁浅空白，图 3/4/5/8 在不同查看器分页下同理。处置：五张栏宽图
  全部改为"栏顶浮动"（同一浮动表机制，horzAnchor=text 锚定本栏、页顶起排、
  上下环绕），至此全部 8 图均浮动、正文全程连续填充——与刊印论文/LaTeX 的
  图浮动模型一致。_qa/whitespace_scan.py 升级为整幅+左右分栏三路检测。
  正文文字与图表数据零变化；表格件数 10→15（7 三线表 + 8 浮动容器）。
  v1–v7 保留）
  （v9：浮动图顶底分层 + 图 4/图 5 合带——作者指出 v8 图 6 那页仍有部分
  空白。渲染核验根因：v8 全部浮动图都抢"顶位"，图 4/图 5 与图 6 的锚位
  同落一页时三块互相竞争（图 3 与图 2 同理），引擎按 tblOverlap=never
  互挤或叠置，正文被挤走后留下空白带/题注重叠。处置：①通栏图（图 1/2/6）
  保持页顶浮排（[t] 位），栏宽图（图 3/7/8）改为栏底浮排（[b] 位，
  tblpYSpec=bottom、topFromText 留距），顶底分层互不争位、正文夹排其间；
  ②消融区最拥挤的图 4/图 5 并排合装为一条通栏底浮图带（两图各 8.2 cm，
  合宽 16.4 cm ≤ 版心 16.8 cm，各携自身双语题注），拥挤页至多"页顶一通栏
  图 + 页底一图带"，不再出现三浮动块同页竞争；③表题注两行加 keepNext，
  杜绝题注与表体被栏/页切分（v8 表 5 题注孤悬栏底）；④浮动屏障消除：
  部分引擎（LibreOffice/WPS 类）不把流序在浮动元素之后的正文回填到浮动
  元素之前的页面，浮动块被推后一页时原页栏底便留下空白——将图 1/图 2/
  图带/图 6 的流位各推迟一段（仍先文后图、图不先于引用页），使元素流位
  贴近实际显示页，屏障空白收敛到一行以内。正文文字与图表数据
  零变化；表格件数 15→14（7 三线表 + 6 单图浮动容器 + 1 双图图带）。
  v1–v8 保留）
  （v10：栏宽表浮动容器化——v9 渲染复检发现第 10 页图 4/图 5 底浮图带与
  内嵌表 6 叠印。根因：LibreOffice/WPS 类引擎只对正文文字做浮动环绕，
  不会让内嵌表格的行避让浮动块矩形，表行伸进底浮图带区域时直接叠压；
  且内嵌表与浮动块的相对落位随分页漂移，v9 恰好把两者挤到同页同栏底。
  处置：栏宽三线表 5/6/7 改为"题注+表体+表注"整体装入无边框单格浮动
  容器、栏底浮排（同栏宽图 [b] 槽位，tblOverlap=never），与图浮动块
  走同一互斥推移路径，任何引擎下均不可能叠印；表 1 附近无浮动块保持
  内嵌，表 2/3/4 在单栏夹层区不变。正文文字与图表数据零变化；顶级
  表格件数仍 14（4 内嵌三线表 + 6 单图容器 + 1 图带 + 3 浮动表容器，
  另有 3 张三线表嵌于容器内）。v1–v9 保留）
  （v11：回退 v10 浮动表 + 图带换页 + 表格单元格排版修正——v10 渲染复检
  证明 LibreOffice 类引擎对"浮动表 vs 浮动表"并不执行 tblOverlap=never
  互斥（表 6 容器仍与图 4/图 5 图带叠印，第 11 页），且浮动表容器紧邻
  单双栏分节边界时引发新的整幅空白（第 9 页 5.2 cm，为 v10 新增回归；
  v9 全文无超 3 cm 空白带）。处置：①表 5/6/7 回退为 v9 的内嵌三线表
  （FLOAT_TABLES 清空，机制保留备用）；②叠印改由"错页"根治：图 4/图 5
  通栏底浮图带的流位自 3.4 节（残余 K 分歧段后）后移至 3.5 节首段之后，
  使其落页与内嵌表 6 相隔整页，两者在任何引擎下不再同页竞争（图 4/图 5
  的正文引用仍在其显示页之前或同页，先文后图不变）；③表格换行修正
  （作者意见）：三线表单元格左右边距由模板默认 108 twip 压至 28 twip，
  数值列"0.893±0.004"类条目不再断行为两行；④分组行（表 4/表 7 的
  基准区块行）整行合并单元格，"Anomaly-ShapeNet——40 类…"等组标题
  不再被挤成三四行。正文文字与图表数据零变化。v1–v10 保留）
  （v12：图带回位 + 表 6 提前——v11 渲染复检：表 5/6/7 回退内嵌与单元格
  修正有效（第 9 页整幅空白消除、数值不再断行），但图带后移至 3.5 节
  引发新空洞（锚页第 10 页右栏 12.9 cm + 显示页第 11 页右栏 18.4 cm）：
  LibreOffice 类引擎中，浮动块之后的正文最早只能从该浮动块的"显示页"
  开始回填，锚位与显示页错开一页时锚页从锚点起全部悬空。处置：图 4/
  图 5 图带回到 v9 流位（3.4 节残余 K 分歧段后，锚位显示同页、无屏障
  空洞），表 6 与图带的同页叠印改由表 6 流位提前两段解决（表 6 整体
  在图带矩形上方结束，其间隔约两段正文作分页缓冲）。正文文字与图表
  数据零变化。v1–v11 保留）
  （v13：页眉残线净化 + 图 2 错页 + 图 3/图带分页——云端 QA（node1 soffice
  渲染 + 矢量线坐标对照官方 .doc 模板）新定位三处：①模板 .doc→.docx 转换
  把页眉表格隐藏边框转丢，正稿每页页顶渲染 3×2 网格残线（首页 17 段、续页
  9 段），新增 _sanitize_headers 显式置 nil 并只保留刊头规则线（首页
  thickThinSmallGap 双线、续页 single 细线），逐线对齐官方 .doc 渲染；
  ②表 4 跨页表尾与图 2 页顶通栏浮排同占第 8 页页顶叠印（99 处），图 2 流位
  后移一段错页；③图 3 栏底浮排与图 4/图 5 底浮图带同落第 9 页左栏底叠置
  （LibreOffice 类引擎不执行浮动表互斥，v11 已证），图 2 错页引发的整体
  回流使两者分页，渲染复核确认。输出文件名 v1→v2，v1 保留。正文文字与
  图表数据零变化。v1–v12 保留）

用法： python3 build_txxb_docx.py
"""
import copy
import os
import re
import sys
from collections import OrderedDict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
TEMPLATE = os.path.join(BASE, 'template_extract', 'txxb_tpl_submit.docx')
DRAFT = os.path.join(ROOT, 'DRAFT_v1_zh.md')
BIB = os.path.join(ROOT, 'references.bib')
FIGDIR = os.path.join(ROOT, 'figures')
OUT_FULL = os.path.join(BASE, 'GCAS_图学学报_排版初稿_原稿_v2_20260831.docx')
OUT_BLIND = os.path.join(BASE, 'GCAS_图学学报_排版初稿_盲审稿_v2_20260831.docx')

# 页面几何（由模板 sectPr 实测）：A4，正文宽 9694 twip = 17.13 cm；
# 双栏栏宽 (9694-236)/2 = 4729 twip = 8.34 cm
FULL_W = Cm(16.6)      # 通栏图/表宽度（TXXB 版心 17.0 cm，留余量）
COL_W = Cm(7.9)        # 栏宽图宽度（TXXB 栏宽 8.11 cm）
COL_W_TW = 4600        # 栏宽（twip）：(9638-438)/2，用于栏内表格
FULL_W_TW = 9638       # 全宽（twip）：11906-2×1134

SONG = '宋体'
HEI = '黑体'
KAI = '楷体'
TNR = 'Times New Roman'

# ---------------------------------------------------------------- 参考文献

LATEX_MAP = {
    r'{\"o}': 'ö', r'{\"u}': 'ü', r'{\"a}': 'ä', r"{\'e}": 'é',
    r'{\v c}': 'č', r'{\v s}': 'š', r'{\AE}': 'Æ', r'{\ae}': 'æ',
    r'\&': '&',
}


def latex_clean(s):
    for k, v in LATEX_MAP.items():
        s = s.replace(k, v)
    s = re.sub(r'\{\\v\s+(\w)\}', lambda m: m.group(1) + '\u030c', s)
    s = s.replace('{', '').replace('}', '')
    s = s.replace('---', '—').replace('--', '-')
    return re.sub(r'\s+', ' ', s).strip()


def parse_bib(path):
    """极简 bib 解析：够用于本 bib（字段一行一个，值以 {} 包裹）。"""
    txt = open(path, encoding='utf-8').read()
    entries = OrderedDict()
    for m in re.finditer(r'@(\w+)\{([^,]+),(.*?)\n\}', txt, re.S):
        etype, key, body = m.group(1).lower(), m.group(2).strip(), m.group(3)
        fields = {}
        # body 末尾补 '\n'：条目正则的 '\n\}' 吃掉了最后一行的换行，而字段
        # 正则要求行尾换行，否则每条的最后一个字段（本 bib 恰为 year）丢失。
        for fm in re.finditer(r'(\w+)\s*=\s*\{(.*?)\},?\s*\n', body + '\n', re.S):
            fields[fm.group(1).lower()] = fm.group(2).strip()
        entries[key] = (etype, fields)
    return entries


def gbt_authors(author_field):
    """'Liu, J. and Xie, G. ...' -> 'LIU J, XIE G, CHEN R, et al.'"""
    parts = [a.strip() for a in author_field.split(' and ')]
    names = []
    for a in parts:
        a = latex_clean(a)
        if ',' in a:
            last, first = [x.strip() for x in a.split(',', 1)]
        else:
            bits = a.split()
            last, first = bits[-1], ' '.join(bits[:-1])
        initials = re.findall(r'[A-ZÄÖÜÆČŠ]', first, re.U)
        names.append((last.upper() + ' ' + ' '.join(initials)).strip())
    if len(names) > 3:
        return ', '.join(names[:3]) + ', et al'
    return ', '.join(names)


# 会议缩写 → (会议录全称, '出版地: 出版者')。GB/T 7714—2015 要求会议论文按
# “[C]//会议录全称. 出版地: 出版者, 年: 页码” 著录。键覆盖 references.bib
# 实际出现的全部 18 种 booktitle（构建时扫描核对）；未覆盖的 booktitle 在
# gbt_format 中保留原缩写并以 [S.l.]: [s.n.] 占位，同时计入 missing_report。
PROCEEDINGS_MAP = {
    'CVPR': ('Proceedings of the IEEE/CVF Conference on Computer Vision '
             'and Pattern Recognition', 'Piscataway: IEEE'),
    'CVPR Workshops': ('Proceedings of the IEEE/CVF Conference on Computer '
                       'Vision and Pattern Recognition Workshops',
                       'Piscataway: IEEE'),
    'ICCV': ('Proceedings of the IEEE/CVF International Conference on '
             'Computer Vision', 'Piscataway: IEEE'),
    'WACV': ('Proceedings of the IEEE/CVF Winter Conference on Applications '
             'of Computer Vision', 'Piscataway: IEEE'),
    'ECCV': ('Proceedings of the European Conference on Computer Vision',
             'Cham: Springer'),
    'ACCV': ('Proceedings of the Asian Conference on Computer Vision',
             'Cham: Springer'),
    'NeurIPS': ('Advances in Neural Information Processing Systems',
                'Red Hook: Curran Associates'),
    'NeurIPS Datasets and Benchmarks Track':
        ('Advances in Neural Information Processing Systems',
         'Red Hook: Curran Associates'),
    'AAAI': ('Proceedings of the AAAI Conference on Artificial Intelligence',
             'Palo Alto: AAAI Press'),
    'IJCAI': ('Proceedings of the International Joint Conference on '
              'Artificial Intelligence', '[S.l.]: IJCAI Organization'),
    'ICML': ('Proceedings of the International Conference on Machine '
             'Learning', '[S.l.]: PMLR'),
    'ACM Multimedia': ('Proceedings of the ACM International Conference on '
                       'Multimedia', 'New York: ACM'),
    'SIGGRAPH': ('Proceedings of the Annual Conference on Computer Graphics '
                 'and Interactive Techniques (SIGGRAPH)', 'New York: ACM'),
    '3DIM': ('Proceedings of the International Conference on 3-D Digital '
             'Imaging and Modeling', 'Piscataway: IEEE'),
    'IEEE Visualization': ('Proceedings of IEEE Visualization',
                           'Piscataway: IEEE'),
    'ICRA': ('Proceedings of the IEEE International Conference on Robotics '
             'and Automation', 'Piscataway: IEEE'),
    'VISAPP': ('Proceedings of the International Conference on Computer '
               'Vision Theory and Applications (VISAPP)',
               'Setúbal: SciTePress'),
    # PaDiM 所在会议录（ICPR 2020 国际研讨会卷，Springer LNCS 12664）
    'ICPR Workshops': ('Pattern Recognition. ICPR International Workshops '
                       'and Challenges', 'Cham: Springer'),
}


# 电子资源引用日期（GB/T 7714—2015 要求 [EB/OL] 给出引用日期与获取路径）
CITE_DATE = '2026-08-28'

# 图书出版地补全（GB/T 7714 要求出版地；按出版社官方所在地填制）
BOOK_PLACE_MAP = {
    'Wiley': 'Hoboken',
}


def gbt_format(key, etype, f):
    """转 GB/T 7714—2015 顺序编码制条目；返回 (text, missing_fields)。"""
    missing = []
    au = gbt_authors(f.get('author', ''))
    ti = latex_clean(f.get('title', ''))
    year = f.get('year', '')
    pages = latex_clean(f.get('pages', '')) if 'pages' in f else ''
    if etype == 'article':
        jn = latex_clean(f.get('journal', ''))
        vol = f.get('volume', '')
        num = f.get('number', '')
        s = f'{au}. {ti}[J]. {jn}, {year}'
        if vol:
            s += f', {vol}'
            if num:
                s += f'({num})'
        else:
            missing.append('卷期')
        if pages:
            s += f': {pages}'
        elif 'note' in f:
            s += f'. {latex_clean(f["note"])}'
            missing.append('页码（在线先行出版）')
        else:
            missing.append('页码')
        s += '.'
    elif etype == 'inproceedings':
        bt = latex_clean(f.get('booktitle', ''))
        if bt in PROCEEDINGS_MAP:
            venue, place_pub = PROCEEDINGS_MAP[bt]
        else:
            venue, place_pub = bt, '[S.l.]: [s.n.]'
            missing.append('会议录全称/出版地/出版者（[S.l.]: [s.n.] 占位）')
        s = f'{au}. {ti}[C]//{venue}'
        vol = f.get('volume', '')
        num = f.get('number', '')
        if vol:      # 连续编号会议录（如 AAAI）的卷期作为其他题名信息
            s += f': Vol {vol}' + (f'({num})' if num else '')
        s += f'. {place_pub}, {year}'
        if pages:
            s += f': {pages}'
        else:
            missing.append('页码')
        s += '.'
    elif etype == 'book':
        pub = latex_clean(f.get('publisher', ''))
        ed = f.get('edition', '')
        s = f'{au}. {ti}[M]'
        if ed:
            s += f'. {ed} ed'
        place = BOOK_PLACE_MAP.get(pub)
        if place:
            s += f'. {place}: {pub}, {year}.'
        else:
            s += f'. [S.l.]: {pub}, {year}.'
            missing.append('出版地')
    else:  # misc → arXiv 预印本，按电子资源著录（GB/T 7714：(更新日期)[引用日期]. 获取路径）
        how = latex_clean(f.get('howpublished', ''))
        m_arxiv = re.match(r'arXiv:\s*([0-9.]+)', how)
        if m_arxiv:
            s = (f'{au}. {ti}[EB/OL]. ({year})[{CITE_DATE}]. '
                 f'https://arxiv.org/abs/{m_arxiv.group(1)}.')
        else:
            s = f'{au}. {ti}[EB/OL]. ({year})[{CITE_DATE}]. {how}.'
            missing.append('获取路径（非 arXiv，请终核 URL）')
    return s, missing


# ---------------------------------------------------------------- 引用编号

CITE_ORDER = OrderedDict()   # key -> num


def cite_nums(keys):
    nums = []
    for k in keys:
        k = k.strip()
        if k not in CITE_ORDER:
            CITE_ORDER[k] = len(CITE_ORDER) + 1
        nums.append(CITE_ORDER[k])
    return nums


def sup_label(nums):
    """[1,2,3,7] -> '1-3,7'"""
    nums = sorted(set(nums))
    out, i = [], 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        out.append(str(nums[i]) if i == j else f'{nums[i]}-{nums[j]}')
        i = j + 1
    return ','.join(out)


# ---------------------------------------------------------------- md 表格解析

def load_md_tables(path):
    lines = open(path, encoding='utf-8').read().splitlines()
    anchors = {
        'T1': '**表 1.**', 'T2': '**表 T2-1.**', 'T3': '**表 T2-2.**',
        'T4': '**表 T2-3.**', 'T5': '**表 3.**', 'T6': '**表 4.**',
        'T7': '**表 5.**',
    }
    tables = {}
    for tag, anchor in anchors.items():
        idx = next(i for i, l in enumerate(lines) if l.startswith(anchor))
        rows = []
        i = idx + 1
        while i < len(lines) and not lines[i].strip().startswith('|'):
            i += 1
        while i < len(lines) and lines[i].strip().startswith('|'):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            if not all(re.fullmatch(r':?-{2,}:?', c or '---') for c in cells):
                rows.append(cells)
            i += 1
        tables[tag] = rows
    return tables


# ---------------------------------------------------------------- 底层排版助手

def set_font(run, ascii_=TNR, ea=SONG, size=10.5, bold=None, italic=None,
             sup=False, sub=False):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ascii_)
    rFonts.set(qn('w:hAnsi'), ascii_)
    rFonts.set(qn('w:eastAsia'), ea)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if sup:
        run.font.superscript = True
    if sub:
        run.font.subscript = True


TOKEN_RE = re.compile(
    r'(\*\*.+?\*\*|(?<![A-Za-z0-9*])\*[^*\n]+?\*(?![A-Za-z0-9*])'
    r'|⟦.+?⟧|\^\{.+?\}|_\{.+?\})')


def emit_runs(par, text, size=10.5, ea=SONG, ascii_=TNR, bold=None,
              italic=None, cite_bold=True):
    """解析 **粗体**、*强调(楷体/斜体)*、⟦引用键⟧、^{上标}、_{下标} 并写 run。"""
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2])
            set_font(r, ascii_, HEI, size, bold=True, italic=italic)
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            r = par.add_run(tok[1:-1])
            set_font(r, ascii_, KAI, size, bold=bold, italic=True)
        elif tok.startswith('⟦') and tok.endswith('⟧'):
            nums = cite_nums(tok[1:-1].split(','))
            r = par.add_run('[' + sup_label(nums) + ']')
            set_font(r, TNR, SONG, size, bold=cite_bold, sup=True)
        elif tok.startswith('^{'):
            r = par.add_run(tok[2:-1])
            set_font(r, ascii_, ea, size, bold=bold, italic=italic, sup=True)
        elif tok.startswith('_{'):
            r = par.add_run(tok[2:-1])
            set_font(r, ascii_, ea, size, bold=bold, italic=italic, sub=True)
        else:
            r = par.add_run(tok)
            set_font(r, ascii_, ea, size, bold=bold, italic=italic)


def pset(par, align=None, fli_chars=None, before=None, after=None,
         line=None, snap=None, ind_left=None, hanging=None, line_rule='auto',
         ind_chars_lr=None):
    """段落属性：对齐/首行缩进(字符)/段前段后(磅)/行距/网格对齐/悬挂缩进。

    TXXB 版新增：line_rule='exact' 时 line 为固定行距（twip，正文 15.7 磅
    = 314）；ind_chars_lr=n 时左右各缩进 n 字符（摘要/关键词两边缩 2 格）。"""
    if align is not None:
        par.alignment = align
    pPr = par._p.get_or_add_pPr()
    if fli_chars is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(fli_chars * 100))
        ind.set(qn('w:firstLine'), str(int(fli_chars * 210)))
    if ind_chars_lr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:leftChars'), str(ind_chars_lr * 100))
        ind.set(qn('w:rightChars'), str(ind_chars_lr * 100))
    if ind_left is not None or hanging is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        if ind_left is not None:
            ind.set(qn('w:left'), str(ind_left))
        if hanging is not None:
            ind.set(qn('w:hanging'), str(hanging))
    if before is not None or after is not None or line is not None:
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        if before is not None:
            sp.set(qn('w:before'), str(int(before * 20)))
        if after is not None:
            sp.set(qn('w:after'), str(int(after * 20)))
        if line is not None:
            sp.set(qn('w:line'), str(line))
            sp.set(qn('w:lineRule'), line_rule)
    if snap is False:
        el = OxmlElement('w:snapToGrid')
        el.set(qn('w:val'), '0')
        pPr.insert(0, el)


def add_sect_break(doc, sect_clone):
    """追加一个携带 sectPr 的空段（结束此前内容所在的节）。"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(copy.deepcopy(sect_clone))
    pset(p, line=240, snap=False)
    for r in p.runs:
        r.font.size = Pt(2)
    return p


def pagetop_sect(sect_clone):
    """克隆节属性并把节起始方式由 continuous 改为 nextPage（新页顶起排）。

    用于高通栏块（图 6）：跨页高图在部分查看器（WPS/移动端预览）中会被分页
    边界截断，nextPage 保证该节整体从新一页顶部开始、整页容纳。同时删去
    titlePg 与 first 页眉/页脚引用，避免该节首页套用首页脚注区（收稿日期/
    基金项目）的页脚。
    """
    s = copy.deepcopy(sect_clone)
    t = s.find(qn('w:type'))
    if t is None:
        t = OxmlElement('w:type')
        s.insert(0, t)
    t.set(qn('w:val'), 'nextPage')
    tp = s.find(qn('w:titlePg'))
    if tp is not None:
        s.remove(tp)
    for tag in ('w:headerReference', 'w:footerReference'):
        for ref in s.findall(qn(tag)):
            if ref.get(qn('w:type')) == 'first':
                s.remove(ref)
    return s


_BORDER_SIDES_TC = ('top', 'left', 'bottom', 'right', 'insideH', 'insideV')
_BORDER_SIDES_P = ('top', 'left', 'bottom', 'right', 'between', 'bar')


def _fill_border_el(parent, tagname, sides, keep=None):
    """把 parent 下的 <tagname> 边框组重建为全 nil；keep 为
    {side: (val, sz)} 时该边保留指定线型。按 OOXML 规定的边顺序写入。"""
    old = parent.find(qn(tagname))
    if old is not None:
        parent.remove(old)
    borders = OxmlElement(tagname)
    for side in sides:
        el = OxmlElement('w:' + side)
        if keep and side in keep:
            val, sz = keep[side]
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
        else:
            el.set(qn('w:val'), 'nil')
        borders.append(el)
    return borders


def _sanitize_headers(doc):
    """页眉边框净化（v13）。

    根因（云端 QA 实测定位）：图学学报模板 .doc 经 LibreOffice 转 .docx 时，
    页眉表格的“隐藏边框”被转丢——单元格残留 bottom/insideH 实线、单元格内
    段落又以空 <w:pBdr/> 继承 Header 样式的下边框——正稿每页页顶因此渲染出
    3×2 网格残线（原版 .doc 首页仅一条 thick-thin 双线、续页仅一条细线，
    对照 render_tpl/tpl.pdf 与 render_yg_v1f 第 1/2 页矢量线坐标）。

    处置：对 header1/2/3 的页眉表格——表级与单元格级边框全部显式置 nil，
    仅末行单元格保留刊头规则线（首页 thickThinSmallGap 24/8 pt 双线、续页
    single 6/8 pt 细线）；表内与表后所有段落的 pBdr 全边显式 none，压住
    Header 样式继承的下边框。渲染结果与官方 .doc 逐线一致。"""
    for part in doc.part.package.iter_parts():
        if not re.search(r'header\d+\.xml$', str(part.partname)):
            continue
        hdr = part.element
        texts = ''.join(t.text or '' for t in hdr.iter(qn('w:t')))
        first_page = '图 学 学 报' in texts or 'JOURNAL OF GRAPHICS' in texts
        rule = {'bottom': ('thickThinSmallGap', '24')} if first_page \
            else {'bottom': ('single', '6')}
        for tbl in hdr.iter(qn('w:tbl')):
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is not None:
                borders = _fill_border_el(tblPr, 'w:tblBorders',
                                          _BORDER_SIDES_TC)
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is not None:
                    tblW.addnext(borders)
                else:
                    tblPr.append(borders)
            rows = tbl.findall(qn('w:tr'))
            for ri, tr in enumerate(rows):
                last = (ri == len(rows) - 1)
                for tc in tr.findall(qn('w:tc')):
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        continue
                    borders = _fill_border_el(
                        tcPr, 'w:tcBorders', _BORDER_SIDES_TC,
                        keep=rule if last else None)
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcW.addnext(borders)
                    else:
                        tcPr.insert(0, borders)
        for p in hdr.iter(qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            borders = _fill_border_el(pPr, 'w:pBdr', _BORDER_SIDES_P)
            for side in borders:
                side.set(qn('w:val'), 'none')
                side.set(qn('w:sz'), '0')
                side.set(qn('w:space'), '0')
            pstyle = pPr.find(qn('w:pStyle'))
            if pstyle is not None:
                pstyle.addnext(borders)
            else:
                pPr.insert(0, borders)


def body_para(doc, text, size=10.5, indent=True):
    # TXXB 模板要求 18：正文 5 号宋体，行距固定 15.7 磅（=314 twip）
    p = doc.add_paragraph()
    pset(p, snap=False, line=314, line_rule='exact')
    if indent:
        pset(p, fli_chars=2)
    emit_runs(p, text, size=size)
    return p


def heading1(doc, text):
    # TXXB 模板要求 15：一级标题 4 号黑体，占二行（以段前段后近似）
    p = doc.add_paragraph()
    pset(p, before=13, after=6.5, snap=False, line=300)
    emit_runs(p, text, size=14, ea=HEI, ascii_=HEI)
    return p


def heading2(doc, text):
    # TXXB 模板要求 16：二级标题 5 号黑体，占一行
    p = doc.add_paragraph()
    pset(p, before=6.5, after=3, snap=False, line=280)
    emit_runs(p, text, size=10.5, ea=HEI, ascii_=HEI)
    return p


def heading3(doc, text):
    # TXXB 模板要求 17：三级标题 5 号楷体，占一行
    p = doc.add_paragraph()
    pset(p, before=3, after=0, snap=False, line=280)
    emit_runs(p, text, size=10.5, ea=KAI, ascii_=KAI)
    return p


EQN_COUNT = [0]


def equation(doc, eq_text, num):
    EQN_COUNT[0] += 1
    p = doc.add_paragraph()
    pset(p, snap=False, before=3, after=3, line=260)
    ts = p.paragraph_format.tab_stops
    ts.add_tab_stop(Cm(3.85), WD_TAB_ALIGNMENT.CENTER)
    ts.add_tab_stop(Cm(8.11), WD_TAB_ALIGNMENT.RIGHT)   # 编号贴双栏栏缘
    r = p.add_run('\t')
    set_font(r, TNR, SONG, 10.5)
    emit_runs(p, eq_text, size=10.5)
    r = p.add_run('\t(' + str(num) + ')')
    set_font(r, TNR, SONG, 10.5)
    return p


def keep_next(par):
    """段落 keepNext：与下一块同页/同栏（题注→表体），防题注孤悬栏底。"""
    pPr = par._p.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is None:
        pPr.insert(0, OxmlElement('w:keepNext'))


def caption_pair(doc, zh, en, size=9):
    """表题注：中文行小五黑体、英文行小五宋体/Times（模板体例）。
    两行均置 keepNext 与其后表体绑定（v9，见文件头注记）。
    TXXB 体例：英文行 TNR 加粗（模板样例「Table 1 …」加粗）。"""
    p1 = doc.add_paragraph()
    pset(p1, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=3, after=0,
         line=240)
    keep_next(p1)
    emit_runs(p1, zh, size=size, ea=HEI)
    p2 = doc.add_paragraph()
    pset(p2, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=0, after=3,
         line=240)
    keep_next(p2)
    emit_runs(p2, en, size=size, ea=SONG, bold=True)
    return p1, p2


def table_note(doc, text, size=8):
    # TXXB 表内文字 6 号（8 磅），表注随表内字号
    p = doc.add_paragraph()
    pset(p, snap=False, before=1, after=6, line=220)
    emit_runs(p, '注：' + text, size=size)
    return p


def merge_caption(cap, sub, en=False):
    """分图/图例说明并入题注行（TXXB 模板样例体例，如「图2 帧间差值曲线
    ((a) 首次帧差的曲线图；(b) 二次帧差的曲线图)」）。说明型 sub 自带外层
    括号直接拼接；(a)(b) 分图型补一层半角外层括号（模板即双括号样式）。"""
    if not sub:
        return cap
    if en:
        if sub.startswith('(a)'):
            return cap + ' (' + sub + ')'    # 分图型补外层括号
        return cap + ' ' + sub               # 说明型已带括号
    if sub.startswith('（a）'):
        half = sub.replace('（', '(').replace('）', ')')
        return cap + '(' + half + ')'        # 分图型转半角并补外层括号
    return cap + sub                          # 说明型已带外层括号


def add_figure(doc, png, width, zh_cap, zh_sub, en_cap, en_sub):
    """图 + 双语题注（TXXB 体例）：图下居中两行——中文行小五黑体、英文行
    小五 TNR 加粗；分图 (a)(b) 与图例说明并入题注行。"""
    p = doc.add_paragraph()
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=6, after=2,
         line=240)
    run = p.add_run()
    run.add_picture(os.path.join(FIGDIR, png), width=width)
    lines = [(merge_caption(zh_cap, zh_sub), HEI, None),
             (merge_caption(en_cap, en_sub, en=True), SONG, True)]
    for i, (text, ea, bold) in enumerate(lines):
        p = doc.add_paragraph()
        pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False,
             before=3 if i == 0 else 0,
             after=3 if i == len(lines) - 1 else 0, line=240)
        emit_runs(p, text, size=9, ea=ea, bold=bold)


def _float_props(tbl, box_tw, horz_anchor, y_spec, top_gap, bot_gap,
                 grid=None):
    """浮动表通用属性：重建 tblpPr/tblOverlap/tblW/tblLayout/tblCellMar，
    保持 OOXML 规定顺序；grid 为各列宽（twip）列表，缺省单列 box_tw。"""
    tblPr = tbl.tblPr
    keep_look = tblPr.find(qn('w:tblLook'))
    for el in list(tblPr):
        tblPr.remove(el)
    # QA 已知怪癖：LibreOffice 6.0 渲染双栏节浮动表时会把单元格内容整体
    # 右移约一个栏距（表框位置正确），OEP v12 同现象、Word/WPS 多轮校审
    # 无此偏移；LO 校样中图片右缘因此视觉上超出版心 0.4–0.7 cm，以 Word
    # 实际打开为准（对照实验 _qa/_float_exp/）。
    tblpPr = OxmlElement('w:tblpPr')
    for k, v in (('leftFromText', '0'), ('rightFromText', '0'),
                 ('topFromText', top_gap), ('bottomFromText', bot_gap),
                 ('vertAnchor', 'margin'), ('horzAnchor', horz_anchor),
                 ('tblpXSpec', 'center'), ('tblpYSpec', y_spec)):
        tblpPr.set(qn('w:' + k), v)
    tblPr.append(tblpPr)
    ov = OxmlElement('w:tblOverlap')
    ov.set(qn('w:val'), 'never')
    tblPr.append(ov)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), box_tw)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tblPr.append(lay)
    mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), '0')
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)
    if keep_look is not None:
        tblPr.append(keep_look)
    gcols = tbl.find(qn('w:tblGrid')).findall(qn('w:gridCol'))
    for gc, w in zip(gcols, grid or [box_tw]):
        gc.set(qn('w:w'), str(w))


def _fill_figure_cell(cell, png, width, zh_cap, zh_sub, en_cap, en_sub):
    """浮动容器单元格内容：图 + 双语题注（TXXB 体例：中文行黑体、英文行
    TNR 加粗，分图/图例说明并入题注行）。"""
    p = cell.paragraphs[0]
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=2, after=2,
         line=240)
    run = p.add_run()
    run.add_picture(os.path.join(FIGDIR, png), width=width)
    lines = [(merge_caption(zh_cap, zh_sub), HEI, None),
             (merge_caption(en_cap, en_sub, en=True), SONG, True)]
    for i, (text, ea, bold) in enumerate(lines):
        p = cell.add_paragraph()
        pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False,
             before=3 if i == 0 else 0,
             after=2 if i == len(lines) - 1 else 0, line=240)
        emit_runs(p, text, size=9, ea=ea, bold=bold)


def _pad_anchor(doc):
    """浮动表后垫一个 2 pt 空锚段：部分引擎（LibreOffice 类）把浮动框与其
    相邻段落绑页，浮动框被推到下页时相邻正文段会陪跳、在原页留下空白；
    空锚段代为"陪跳"，正文段落留在原页正常回填。Word/WPS 下不可见、无害。"""
    anchor = doc.add_paragraph()
    pset(anchor, snap=False, before=0, after=0, line=240)
    r = anchor.add_run('')
    r.font.size = Pt(2)


def add_floating_figure(doc, png, width, zh_cap, zh_sub, en_cap, en_sub,
                        column=False, y_override=None):
    """浮动图文框：图 + 双语题注整体装入无边框单格浮动表。

    表以 w:tblpPr 锚定、tblOverlap=never，正文连续环绕。column=False 为
    通栏图（横跨双栏、页顶浮排，即刊印论文的 [t] 浮动图）；column=True
    为栏宽图（horzAnchor=text 锚定本栏、栏底浮排，即 [b] 浮动图）。
    v9 起栏宽图由栏顶改为栏底：栏顶位与通栏图的页顶位同页时互相竞争，
    多引擎下会叠置或互挤出空白（见文件头 v9 注记），顶底分层后通栏图占
    页顶、栏宽图占栏底，正文夹排其间。相比内嵌排版：不可拆图块不再把
    栏底/页底搁浅成空白，且整体成块浮动，不会被分页截断。
    """
    table = doc.add_table(rows=1, cols=1)
    # TXXB 版心：栏宽 4600 twip（8.11 cm）/ 正文全宽 9638 twip（17.0 cm）
    # （修复 fork 遗留的 OEP 硬编码 4729/9694，后者超出 TXXB 版心 0.1 cm）
    box_tw = str(COL_W_TW) if column else str(FULL_W_TW)
    # 顶浮块与正文的间距留在其下缘，底浮块留在其上缘；y_override 允许
    # 通栏图改页底浮排（八轮：图 6 与图带顶底换槽）
    y_spec = y_override or ('bottom' if column else 'top')
    _float_props(table._tbl, box_tw,
                 horz_anchor='text' if column else 'margin',
                 y_spec=y_spec,
                 top_gap='170' if y_spec == 'bottom' else '0',
                 bot_gap='0' if y_spec == 'bottom' else '170')
    _fill_figure_cell(table.cell(0, 0), png, width, zh_cap, zh_sub,
                      en_cap, en_sub)
    _pad_anchor(doc)
    return table


def add_floating_figure_band(doc, fig_a, fig_b, y_spec='bottom'):
    """通栏底浮图带：两张栏宽图并排合装进一个 1×2 无边框浮动表（v9）。

    消融区的图 4/图 5 与图 6 引用位相邻，三块独立浮动时锚位同落一页、
    竞争顶/底位，多引擎下互挤出空白带。合带后两图共占一条页底通栏带
    （各携自身双语题注，合宽 16.4 cm ≤ 版心 16.8 cm），拥挤页至多
    "页顶一通栏图 + 页底一图带"，正文夹排其间。
    """
    table = doc.add_table(rows=1, cols=2)
    half = FULL_W_TW // 2
    _float_props(table._tbl, str(FULL_W_TW), horz_anchor='margin',
                 y_spec=y_spec,
                 top_gap='170' if y_spec == 'bottom' else '0',
                 bot_gap='0' if y_spec == 'bottom' else '170',
                 grid=[half, FULL_W_TW - half])
    _fill_figure_cell(table.cell(0, 0), *fig_a)
    _fill_figure_cell(table.cell(0, 1), *fig_b)
    # 两格内容不等高时垂直居中，格内空白上下分摊（txxb-v1 QA：图 7/图 8
    # 高度差 1.5 cm，顶对齐时矮格上方空白使旁栏正文提前止排）
    for cell in (table.cell(0, 0), table.cell(0, 1)):
        tcPr = cell._tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), 'center')
        tcPr.append(va)
    _pad_anchor(doc)
    return table


# v11：清空——LibreOffice 类引擎不执行浮动表间的 tblOverlap=never 互斥，
# v10 的表容器仍与图带叠印且在分节边界引发新空白；表 5/6/7 回退内嵌，
# 叠印改由图带错页解决（见文件头 v11 注记）。机制保留备用。
FLOAT_TABLES = ()


def add_floating_table(doc, tag, tables):
    """栏宽三线表浮动容器（v10）：题注两行 + 表体 + 表注整体装入无边框
    单格浮动表，栏底浮排（与栏宽图同 [b] 槽位）。

    动机：内嵌三线表的表行不会环绕避让底部浮动块（LibreOffice/WPS 类
    引擎不对表行做浮动环绕换行），表行伸进底浮块矩形时直接叠印（v9
    第 10 页表 6 与图 4/图 5 图带重叠）。改为浮动容器后与其它浮动块
    走 tblOverlap=never 的互斥路径，引擎自动推移、不再叠印；题注/
    表体/表注成块浮动，keepNext 亦不再需要。仅适用于整块可容于一栏
    的栏宽表（浮动块不可拆行）：表 5/6/7 连题注最高约 11 cm，远小于
    栏高 24.5 cm；表 1 附近无浮动块，保持内嵌（可跨页续排）。"""
    zh, en, note, tw, size, ratio = TABLE_META[tag]
    container = doc.add_table(rows=1, cols=1)
    _float_props(container._tbl, str(tw), horz_anchor='text',
                 y_spec='bottom', top_gap='170', bot_gap='0')
    cell = container.cell(0, 0)
    p1 = cell.paragraphs[0]
    pset(p1, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=3, after=0,
         line=240)
    emit_runs(p1, zh, size=9, ea=HEI)
    p2 = cell.add_paragraph()
    pset(p2, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=0, after=3,
         line=240)
    emit_runs(p2, en, size=9, ea=SONG)
    rows = tables[tag]
    group_rows = tuple(i for i, r in enumerate(rows)
                       if r[0].startswith('*') and all(not c for c in r[1:]))
    add_3line_table(cell, rows, tw, size, first_col_ratio=ratio,
                    group_rows=group_rows)
    # cell.add_table 会自动补一个尾段（OOXML 要求单元格以段落收尾），
    # 直接用作表注段，避免多余空段撑高容器
    pn = cell.paragraphs[-1]
    pset(pn, snap=False, before=1, after=2, line=220)
    emit_runs(pn, '注：' + note, size=7.5)
    _pad_anchor(doc)
    return container


def _borders(el_parent, spec):
    """spec: list of (tagname, sz) —— single 线；sz 单位 1/8 pt。"""
    borders = OxmlElement('w:tblBorders' if el_parent.tag == qn('w:tblPr')
                          else 'w:tcBorders')
    for tag, sz in spec:
        b = OxmlElement('w:' + tag)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        borders.append(b)
    el_parent.append(borders)


def add_3line_table(doc, rows, total_tw, size, first_col_ratio=None,
                    group_rows=(), align_first_left=True):
    """三线表：上下 1 pt 粗线 + 表头下 0.5 pt 栏目线，无竖线。"""
    ncols = max(len(r) for r in rows)
    rows = [r + [''] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    _borders(tblPr, [('top', 8), ('bottom', 8)])
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tw = tblPr.find(qn('w:tblW'))
    if tw is None:
        tw = OxmlElement('w:tblW')
        tblPr.append(tw)
    tw.set(qn('w:w'), str(total_tw))
    tw.set(qn('w:type'), 'dxa')
    # v11：单元格左右边距由模板默认 108 twip 压至 28 twip——数值列
    # "0.893±0.004"类条目在 6.5 pt 下不再被边距挤成两行（作者意见）；
    # v12：超宽逐类表（>10 列，表 2/3）进一步压至 8 twip（三线表无竖线，
    # 视觉无影响），配合表头紧排让 "gemstone" 等长类别名单行容纳
    cell_mar = 8 if ncols > 10 else 28
    mar = OxmlElement('w:tblCellMar')
    for side, w in (('top', 0), ('left', cell_mar), ('bottom', 0),
                    ('right', cell_mar)):
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(w))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)
    # 列宽
    if first_col_ratio:
        w0 = int(total_tw * first_col_ratio)
        wrest = (total_tw - w0) // (ncols - 1)
        widths = [w0] + [wrest] * (ncols - 1)
    else:
        widths = [total_tw // ncols] * ncols
    grid = table._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(w))
    # v11：分组行（如表 4/表 7 的基准区块标题行）整行合并单元格，
    # 组标题占满表宽、不再被首列宽度挤成三四行（作者意见）
    for i in group_rows:
        table.cell(i, 0).merge(table.cell(i, ncols - 1))
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if i in group_rows and j > 0:
                continue                     # 已并入合并单元格
            cell = table.cell(i, j)
            tcPr = cell._tc.get_or_add_tcPr()
            # python-docx 0.8.11 不支持 cell.width = None；等价做法：
            # 移除既有 w:tcW 后追加显式 dxa 宽度（合并的组行取全表宽）
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcw = OxmlElement('w:tcW')
            tcw.set(qn('w:w'), str(total_tw if i in group_rows
                                   else widths[j]))
            tcw.set(qn('w:type'), 'dxa')
            tcPr.append(tcw)
            va = OxmlElement('w:vAlign')
            va.set(qn('w:val'), 'center')
            tcPr.append(va)
            if i == 0:
                _borders(tcPr, [('bottom', 4)])
            p = cell.paragraphs[0]
            left = align_first_left and j == 0 and i > 0 and \
                i not in group_rows
            pset(p, align=WD_ALIGN_PARAGRAPH.LEFT if left
                 else WD_ALIGN_PARAGRAPH.CENTER,
                 snap=False, before=1, after=1, line=220)
            txt = cell_text.replace(' ± ', '±').replace('‡', '^{‡}')
            # 千分位统一为 GB/T 15835 分节空格
            txt = re.sub(r'(\d),(\d{3})', '\\1\u2009\\2', txt)
            if i in group_rows and j == 0:
                emit_runs(p, txt, size=size)
            elif txt:
                emit_runs(p, txt, size=size)
                # v12：超宽表表头的长英文类别名（如 gemstone）轻微紧排
                # （-0.2 pt/字符，视觉不可辨），保证单行容纳
                if (i == 0 and j > 0 and ncols > 10 and len(txt) >= 8
                        and txt.isascii()):
                    for r in p.runs:
                        sp = OxmlElement('w:spacing')
                        sp.set(qn('w:val'), '-4')
                        r._r.get_or_add_rPr().append(sp)
    return table


# ---------------------------------------------------------------- 正文内容

ZH_TITLE = '免训练几何一致评分的三维点云异常检测'
EN_TITLE = ('Geometry-consistent anomaly scoring for training-free 3D '
            'point cloud anomaly detection')
SHORT_TITLE = ZH_TITLE

# 摘要按"目的—方法—结果—结论"四要素连贯成段书写，不加显式引导词
# （与《光学 精密工程》已刊文章体例一致，如 2025 年第 33 卷已刊摘要均为
# "针对……提出……。首先……；其次……；最后……。实验结果表明：……"式连续段落）。
ABSTRACT_ZH = (
    '三维点云异常检测需在仅有少量无缺陷模板的条件下识别并定位缺陷，'
    '基于配准的免训练管线可快速适配产线换型，但其评分阶段长期沿用粗粒度设计，'
    '存在正常公差池化过粗、全局配准适配度加权无法抑制局部错配模板、'
    '逐点独立打分忽略缺陷空间连续性三处局限。针对上述问题，'
    '提出免训练的几何一致异常评分（Geometry-Consistent Anomaly Scoring，GCAS）'
    '模块，以三级算子替换常规评分阶段：'
    '首先，细粒度公差校准在小邻域（K = 10）内估计局部正常公差，以保留微小缺陷；'
    '其次，软最小模板融合对取负标准化残差做逐点 softmax，抑制错配模板伪残差并'
    '保留跨模板共识；最后，kNN 图空间一致性正则化在测试扫描的内在邻居图上以无量纲'
    '高斯核巩固连续缺陷响应。实验结果表明：在 Real3D-AD 基准三随机种子、'
    '官方坐标网格协议下，GCAS 将对象级 AUROC 由 0.893 提升至 0.928，点级 AUROC '
    '由 0.896 提升至 0.958，点级 AP 由 0.479 提升至 0.580，对象级 AUROC 超过'
    '已发表最优结果（0.859）；点级增益可迁移至 Anomaly-ShapeNet 与 MulSen-AD，'
    '评分阶段新增耗时不足端到端推理总耗时的 1%。上述结果表明，仅改进评分阶段'
    '即可在零训练成本下显著提升'
    '基于配准的三维异常检测性能，适用于少模板、快速换型的高精度工业质检场景。'
)

KEYWORDS_ZH = '三维异常检测；点云；工业质检；点云配准；异常评分；空间一致性正则化'

# TXXB 要求：英文摘要 ≥2 500 字符、避免 "this paper"；按已刊文章惯例
# （DOI …2026010029/…2026020380）采用被动语态与过去时叙述
ABSTRACT_EN = (
    'Industrial inspection of high-resolution point clouds must identify and '
    'localize surface defects using only a handful of defect-free template '
    'scans, and registration-based pipelines are attractive in this setting '
    'since they require no training and adapt quickly to frequent product '
    'changeover. However, the scoring stage that converts per-template '
    'residuals into anomaly scores has received little attention: normal '
    'tolerances are pooled over neighborhoods far larger than a small '
    'defect, fusion weighted by global registration fitness cannot suppress '
    'a locally misaligned template, and independent point-wise scoring '
    'ignores the spatial contiguity of real defects. To address these three '
    'limitations, a training-free scoring module termed geometry-consistent '
    'anomaly scoring (GCAS) was proposed, which replaced the conventional '
    'calibration-fusion-aggregation stage while keeping registration and '
    'residual extraction unchanged. First, a fine-grained tolerance '
    'calibration stage estimated the local normal tolerances over small '
    'neighborhoods (K = 10), so that small defects were no longer absorbed '
    'into coarsely pooled statistics. Then, a soft-minimum template fusion '
    'stage applied a point-wise softmax over negated standardized residuals, '
    'which suppressed the spurious residuals produced by locally misaligned '
    'templates while preserving cross-template consensus; a formal '
    'proposition further bounded the influence of any single contaminated '
    'template on the fused score. Finally, a kNN-graph spatial consistency '
    'regularization stage consolidated contiguous defect responses over the '
    'intrinsic neighbor graph of the test scan with a dimensionless Gaussian '
    'kernel, so that isolated noise spikes were attenuated while connected '
    'defect regions were reinforced. All three stages reused quantities '
    'already computed by the baseline pipeline, and the complete scoring '
    'module ran in about 21 ms per scan on a single CPU core, adding less '
    'than 1% of the end-to-end inference time. Experimental results '
    'demonstrated that, under a strict protocol on the Real3D-AD benchmark '
    '(three random seeds, the official coordinate grid, and one '
    'benchmark-wide parameter setting), GCAS improved the object-level area '
    'under the receiver operating characteristic curve (AUROC) from 0.893 '
    'to 0.928, the point-level AUROC from 0.896 to 0.958, and the '
    'point-level average precision (AP) from 0.479 to 0.580 over the '
    'identical pipeline with the conventional scoring stage, and the '
    'object-level result exceeded the strongest published one (0.859). On '
    'Anomaly-ShapeNet and MulSen-AD, the point-level gains transferred '
    'directly under the frozen setting, and a label-free parameter '
    're-selection protocol using only the target-benchmark templates '
    'restored and surpassed the object-level performance (from 0.832 to '
    '0.926 on Anomaly-ShapeNet). These results indicate that upgrading only '
    'the scoring stage yields significant gains at zero training cost, '
    'making registration-based 3D anomaly detection more practical for '
    'few-template industrial inspection with frequent product changeover.'
)

KEYWORDS_EN = ('3D anomaly detection; point cloud; industrial inspection; '
               'point cloud registration; anomaly scoring; spatial '
               'consistency regularization')

INTRO = [
    '工业 3D 异常检测旨在高分辨率点云上识别缺陷物体并定位缺陷区域，是高精度制造质检的'
    '关键环节。由于缺陷稀少、形态多样且标注昂贵，该任务通常在无监督设定下开展：训练时'
    '仅有无缺陷数据可用。在此设定下，正常数据本身同样稀缺：Real3D-AD 基准⟦Reg3D-AD⟧'
    '每类仅提供 4 个无缺陷模板扫描——该基准全部点云均由 PMAX-S130 双目蓝光结构光'
    '工业扫描仪对实物工件实测采集（点精度 0.011–0.015 mm，点间距 0.04–0.07 mm，'
    '360°多视角无盲区覆盖），并非仿真生成——这真实反映了产线上采集经标定高分辨率'
    '扫描的实际成本。'
    '如何基于如此少量的正常模板完成缺陷检测与定位，仍是一项极具挑战性的任务。',

    '现有方法大体分为基于学习与基于配准两类。基于学习的方法或通过重建刻画正常性，'
    '如复原被掩蔽的正常几何（IMRNet⟦IMRNet⟧）、对伪异常变体去噪（R3D-AD⟦R3D-AD⟧、'
    'PO3AD⟦PO3AD⟧）；或通过特征嵌入，将测试特征与预训练特征记忆库匹配（M3DM⟦M3DM⟧、'
    'ISMP⟦ISMP⟧、Group3AD⟦Group3AD⟧）。这类方法基准性能优异，但需要针对每个类别进行'
    '数小时的训练，且每次产线换型都须重新训练，在少模板场景下并不实用。相比之下，'
    '基于配准的方法⟦Reg3D-AD,PointCore,3DKeyAD⟧将测试扫描对齐到正常模板、计算逐点几何'
    '残差，并标记残差超出局部正常公差的区域；此类方法无需训练、几分钟即可适配新类别，'
    '且全流程具备可解释性。然而，与研究充分的配准和特征表示环节相比，此类管线中负责'
    '残差标准化、逐模板证据融合以及点分数聚合的评分阶段，所受到的关注要少得多。',

    '在基于学习的路线中，重建类方法与二维检测中的伪异常增广路线一脉相承'
    '⟦DRAEM,CutPaste⟧，PO3AD⟦PO3AD⟧进一步在法向引导增广下预测逐点偏移以强化伪异常'
    '监督；特征嵌入类方法将二维检测的记忆库范式⟦PatchCore,SPADE⟧（在稠密图像基准'
    '⟦MVTecAD⟧上发展）移植到点云：M3DM⟦M3DM⟧融合在 ShapeNet⟦ShapeNet⟧上预训练的'
    '冻结点 Transformer 特征⟦PointTransformer,PointMAE,PointBERT⟧与图像特征，ISMP'
    '⟦ISMP⟧以内部伪模态投影加以丰富，Group3AD⟦Group3AD⟧引入组级对比结构。近期工作'
    '还包括姿态不变形状表示（PASDF⟦PASDF⟧）、配准目标诱导的旋转不变特征（Reg2Inv'
    '⟦Reg2Inv⟧）、曲率增强的自监督表示（CASL⟦CASL⟧）、保组中心重建（DUS-Net'
    '⟦DUSNet⟧）、判别式符号距离函数（DLF-3AD⟦DLF3AD⟧）、带边界感知细化的原型记忆'
    '（M2P-AD⟦M2PAD⟧）、统一多类别建模（MC3D-AD⟦MC3DAD⟧、SeDiR⟦SeDiR⟧）与手工'
    '多尺度描述符（Simple3D⟦Simple3D⟧）等方向；带几何描述子的师生蒸馏⟦3DST,AST⟧、'
    '轻量架构⟦EasyNet⟧与视图投影混合⟦CPMF⟧面向互补的深度扫描基准'
    '⟦MVTec3D,Eyecandies⟧，更广泛的进展见综述⟦ADSurvey⟧。上述方法多数需按类别开展'
    '数小时 GPU 训练且每次换型都须重训，部分还依赖大型冻结骨干网络，与本文所面向的'
    '仅 4 个模板、零训练场景互补。',

    '在基于配准与模板的路线中，Reg3D-AD⟦Reg3D-AD⟧将 RANSAC/ICP 配准⟦RANSAC,ICP⟧与'
    '坐标和冻结特征的双记忆库配对，是标准的免训练基线；PointCore⟦PointCore⟧以单一'
    '联合局部—全局记忆库、点到面 ICP 精化⟦P2LICP⟧与基于秩的分数归一化对其加速；'
    'Template3D-AD⟦Template3DAD⟧用基于曲率的局部特征把中心点与单个对齐模板匹配，'
    '3DKeyAD⟦3DKeyAD⟧配准多个原型并汇入关键点引导的参考集；FPFH⟦FPFH⟧等经典手工'
    '描述子与原始几何残差⟦BTF⟧是若干此类系统的底层。这些管线的共同点在于：评分均'
    '相对于合并后的单一参考基准（记忆库或合并原型）进行，模板间的分歧在评分前已被抹去，公差'
    '要么是全局性的、要么是隐式的。在评分侧，二维检测中逐位置正常统计建模（PaDiM'
    '⟦PaDiM⟧）与分数图高斯平滑（PatchCore⟦PatchCore⟧）均已是成熟做法；而三维点云'
    '缺少规范网格，扫描的内在邻居图⟦GSP,Taubin,BilateralMesh⟧很少被利用，现有点云'
    '方法要么逐点独立打分，要么依赖网络感受野⟦PointNet,PointNetPP⟧获得隐式平滑；'
    '中位数、截尾、软最小等鲁棒统计是估计中的标准工具⟦Huber⟧，鲁棒选择亦被用于含噪'
    '二维记忆库⟦SoftPatch⟧，但其在多模板残差融合上的系统应用——尤其是全局配准'
    '适配度（fitness）加权的失效模式及其逐点修正——在三维异常检测中尚未被深入探讨。',

    '本文发现，常规评分阶段存在三处局限。（1）公差统计池化过粗。正常公差由跨模板残差'
    '在固定空间邻域内池化估计（在本文所基于的常规阶段中为 50 个最近邻），隐含假设缺陷'
    '大于池化半径；仅影响少数点的缺陷会抬高自身所在处的局部公差，在标准化后很大程度上'
    '被自身抬高的公差所抵消。（2）全局适配度无法否决局部配准失败。逐模板残差按 fitness '
    '平方加权融合，预设配准质量可用单个逐模板标量刻画；但实践中归一化后的权重近乎均匀'
    '（在全部 1 206 个测试扫描与三个种子上，四个模板权重的最大值平均仅 0.29，仅略高于'
    '均匀值 0.25），对齐不良的模板因此几乎未受衰减地将伪残差引入融合分数。（3）独立点'
    '打分忽略缺陷连续性。真实表面缺陷占据空间连续的区域，残差噪声则以孤立点为主；逐点'
    '独立打分会增多假阳性并稀释真实缺陷区域，其代价集中体现在点级指标上，而点级恰是基于'
    '配准的方法历来落后于基于学习方法之处。',

    '针对上述局限，本文提出几何一致异常评分（GCAS）：一个免训练的评分模块，用于'
    '替换常规的“校准—融合—聚合”阶段，同时保持配准与残差提取环节不变。第 1 级细粒度'
    '公差校准在小邻域上池化跨模板残差统计，使小缺陷在标准化后得以有效保留；第 2 级'
    '软最小模板融合对取负的标准化残差做逐点 softmax，在模板间存在分歧之处抑制来自'
    '错配模板的单侧离群值，同时保留真实缺陷在每个对齐下均可见所形成的共识性响应；'
    '第 3 级 kNN（k 近邻）图空间一致性正则化以无量纲高斯核在测试扫描的内在邻居图上'
    '平滑融合分数，削弱孤立响应、巩固连续区域。三级均免训练，对端到端推理时间的增量'
    '不足 1%，且仅复用管线中已有的计算结果。',

    '本文的主要贡献如下：（1）指出了基于配准的 3D 异常检测中常规评分阶段存在的三处'
    '局限，并提出免训练、即插即用的评分模块 GCAS，以可忽略的额外推理开销逐一解决上述'
    '问题；（2）提出细粒度公差校准、软最小模板融合与 kNN 图空间一致性正则化三级设计，'
    '逐级消融、参数敏感性分析与负对照实验分别验证了各级的独立贡献；（3）采用“可复现'
    '优先”的实验协议：三个独立种子并报告均值与标准差、全基准单一参数设定且不做逐类'
    '调参、点级评分采用官方坐标网格，全部脚本、冻结参数与逐样本输出通过公开仓库提供'
    '（见文末数据与程序可用性说明）；（4）在 '
    'Real3D-AD、Anomaly-ShapeNet 与 MulSen-AD 上的大量实验证明了 GCAS 的有效性：其在 '
    'Real3D-AD 上取得平均对象级 AUROC 0.928 与点级 AUROC 0.958，对象级超过最强已发表'
    '条目（0.859）。',
]

S21 = [
    '对每个类别，给定 M 个无缺陷模板扫描 {T_{1}, …, T_{M}}（Real3D-AD 上 M = 4），'
    '需在两个层级为测试扫描 S 打分：判定整体是否异常的对象级分数，以及定位异常区域的'
    '点级分数。设置中既无异常数据或标签，亦无额外的正常扫描；所有组件均无需训练。',

    '如图 1 所示，本文管线遵循既有的基于配准的骨架。每个扫描确定性降采样到 '
    'N = 10 000 点；逐点法向与曲率由 k_{nn} = 50 最近邻的局部协方差分析估计'
    '⟦Hoppe,Pauly⟧，并保留测试扫描的 k-NN 索引与距离表供后续阶段复用。测试扫描独立'
    '配准到每个模板：FPFH 对应⟦FPFH⟧上的多起点 RANSAC⟦RANSAC⟧，随后点到点 ICP'
    '⟦ICP,ICPVariants⟧（每模板 8 次重启；保留 fitness 最高的重启）。对每个模板 t 提取'
    '三个逐点残差通道：最近邻（Chamfer）距离 r^{cd}、法向偏差 r^{nm}、曲率差 r^{cv}，'
    '各为 R^{M×N} 中的矩阵，其第 t 行在对齐到模板 t 下计算，q_{t}(x) 表示测试点 x 的'
    '最近模板点。',

    '随后进入常规评分阶段：(a) 基于在 K = 50 个模板邻域点上池化得到的局部正常公差对'
    '残差进行标准化，(b) 以固定权重 (1, 0.5, 0.3) 线性组合通道，(c) 按 fitness 平方'
    '加权平均融合 M 个模板行，(d) 以标准化 top-k 均值（k = 100）把点分数聚合为对象'
    '分数。GCAS 替换 (a) 与 (c)，并在 (c) 与 (d) 之间插入一个空间正则化级；通道组合 '
    '(b) 与聚合 (d) 则与基线完全保持一致，从而确保全部改进均归因于这三级设计。'
    '最终评分阶段的端到端流程与复杂度分析见第 1.6 节。',
]

S22 = [
    '原始残差在整个表面上并不可比：薄壁结构、高曲率区域与点密度变化都会使'
    '无缺陷几何产生系统性的残差水平差异，补救办法是建立局部零模型（null model）。'
    '对全部有序模板对做配准，即可在每个模板位置采集“无缺陷几何在此处产生的残差”样本。'
    '沿用基线管线，该样本由全部 M^{2} 个有序对构建（M = 4 时共 16 对），其中含 M 个'
    '自对；自对残差处于数值噪声量级，相当于收紧池化公差的零值条目。基线与 GCAS 共享'
    '同一构建，对照因此不受影响；排除自对仅使两种方法的绝对指标数值变动小于 0.02，'
    '不改变任何结论（补充材料 C）。',

    '公差表的构建方式如下：固定参考模板 t，每完成一次模板对配准，就经最近邻对应将'
    '残差归集到 t 上的对应模板点。对模板'
    '点 p，令 A_{t}^{(c)}(p) 表示归集到 p 在 t 上 K 个最近邻处的通道 c 残差多重集；'
    '池化规模 K 的选择见下文。公差统计为',
    '@EQ1',
    '@EQ2',
    '其中 σ 以 ρ·|μ| + ε₀（ρ = 0.3，ε₀ = 10^{−6}）为下限，以保证数值稳定。测试残差按',
    '@EQ3',
    '标准化，其中 ε = 10^{−8}；池化样本为空的模板位置回退为 (μ, σ) = (0, 1)。随后将'
    '各通道组合为逐模板证据',
    '@EQ4',
    '通道权重 (λ_{cd}, λ_{nm}, λ_{cv}) = (1, 0.5, 0.3) 原样继承自基线管线；对其重新'
    '调参的做法已作为负对照考察并否决'
    '（补充材料 D）。',

    '在池化规模的选择上，常规阶段取 K = 50，其池化区域远大于小缺陷的尺度，'
    '小缺陷因此抬高自身所在处的公差，其响应在式 (3) 的标准化中被部分抵消。'
    '本文将池化规模缩小到 K = 10。该改动复用'
    '已存储的邻居表、无额外开销，并使精细结构周围的公差估计更为精准。K 决定了偏差—'
    '方差权衡：K 过小时 μ、σ 仅由极少残差样本估计、噪声偏大，K = 5 已使对象级 AUROC '
    '退化；K 过大则过度平滑（第 2.4 节）。K 在全基准上一次固定，从不按类别调参'
    '（第 1.5 节）。',
]

S23 = [
    'M 个标准化行 z_{1}(x), …, z_{M}(x) 须归并为每点一个分数。常规融合按配准 '
    'fitness 的平方为各行加权，但该方案存在两个缺陷。其一，归一化后的 fitness 权重在'
    '实践中近乎均匀（第 2.4 节），无法有效抑制配准不良的模板；其二，更根本的是，配准'
    '错误造成的影响是逐点而非逐模板的：在局部结构附近对齐失败的模板仅在该处产生伪'
    '残差，在其余位置仍具信息量，因此任何逐模板的标量加权在粒度上均不够精确。此外，'
    '噪声污染是单侧的——错配只会抬高残差——合适的算子应在每一点上对高侧离群值保持'
    '鲁棒。',

    '为此，本文对取负的标准化残差做逐点 softmax（玻尔兹曼加权平均⟦Boltzmann⟧）融合：',
    '@EQ5',
    '@EQ6',
    '温度 β ≥ 0。式 (5)、(6) 定义了一个连续族，其端点是两种经典融合：',
    '@EQ7',
    '@EQ8',
    '即 β 在跨模板均值与逐点最小值之间插值。取中间值的 β 会将权重从“相较于同一点其他'
    '模板异常偏高”的行中转移开，而这种偏高现象正是局部配准错误的典型表征；而在所有'
    '模板表现一致的点处（包括在各对齐状态下残差均升高的真实缺陷点），各模板权重则基本'
    '保持均衡。这种不对称性正是区分配准失败与真实缺陷的关键，且无需显式估计哪个模板在'
    '何处失败。',

    '上述融合算子满足三条基本保证（证明见补充材料 I）。',

    '**命题 1.**令 s(x) = Σ_{t} w_{t}(x) z_{t}(x)，权重由式 (5) 给出，β > 0。记 '
    'z = (z_{1}(x), …, z_{M}(x))，则：(i) 共识保真：若 z_{1} = … = z_{M} = c，则 '
    's(x) = c，即所有模板一致给出的响应证据——包括在每个对齐下均可见的真实缺陷——'
    '能够无衰减地保留；(ii) 插值：min_{t} z_{t} ≤ s(x) ≤ (1/M) Σ_{t} z_{t}，且 '
    's(x) 对 β 单调不增（∂s/∂β = −Var_{w}(z) ≤ 0），随 β 增大从式 (7) 的均值端'
    '（β → 0）单调下降到式 (8) 的最小值端（β → ∞）；(iii) 单模板污染上界：对每个模板下标 j '
    '与每个 z，s(x) ≤ min_{i≠j} z_{i} + (M−1)/(eβ)，其中 e 为自然常数。',

    '第 (iii) 条给出了单侧鲁棒性的定量形式：无论错配模板的伪残差多大，融合分数至多比'
    '其余模板的下限高 (M−1)/(eβ)，在 M = 4、β = 1 时约为 1.10 个标准化单位。常规加权'
    '均值不存在这样的界：在实践中观测到的近乎均匀权重下（见第 2.4 节），单行抬高 Δ 会'
    '使融合分数增加约 Δ/M，随 Δ 无界增长。',

    '反观各替代方案：硬最小值（β → ∞）对离群值的抑制最强，但丢弃了跨模板一致性，'
    '对点级指标造成了明显的性能损失（第 2.4 节）；截断或阈值化的 fitness 加权仍受限于'
    '逐模板标量的粗粒度；显式按配准可靠性线索（互最近邻率、反向距离、对应重数）加权'
    '则会主动抑制真实异常——异常区域恰恰是几何结构退化之处，可靠性与异常性在此发生'
    '混淆，本文将其作为负对照报告（补充材料 D）。软最小仅需一个标量 β，且全基准一次'
    '固定（第 1.5 节）。',
]

S24 = [
    '融合之后的分数仍是逐点独立的，而真实缺陷占据空间连续的表面区域、残差噪声'
    '以孤立点为主。基于测试扫描的内在几何结构进行平滑，即可将这一空间先验有效引入评分'
    '中。',

    '具体地，令 N(x) 为 x 在测试扫描上的 k_{nn} 个最近邻（k_{nn} = 50，含自身，复用'
    '预处理表），d(x, y) 为欧氏距离。用逐扫描带宽',
    '@EQ9',
    '正则化分数为核加权邻域平均，即扫描 kNN 图上一步高斯扩散⟦GSP⟧：',
    '@EQ10',
    '将 σ_{S} 与该扫描的中位数邻域距离相绑定，使算子在不同类别与扫描点密度间保持无'
    '量纲；α 是唯一的自由参数。孤立尖峰会被周围低分邻域平滑削弱；连续缺陷区域因相邻点'
    '分数同样升高而得以保留，并在噪声被抑制的背景上展现出更高的对比度。对象级 top-k 统计'
    '量由此能够稳定捕获强化后的缺陷区域，而非噪声尖峰。',

    '平滑强度与对象级聚合相互制约：实验表明，更强的平滑（如迭代扩散）虽能单调改善'
    '点级 AUROC，但会压平分数分布的上尾区，削弱对象级 top-k 统计量的判别力'
    '（第 2.4 节）。因此本文让'
    '两个指标共用一个 α，而非按指标各设强度，并报告权衡曲线。',
]

S25 = [
    'GCAS 引入三个标量：池化规模 K、温度 β 与平滑强度 α。三者均按基准一次冻结，从不按'
    '类别调整、从不接触测试标签，由建立在留一模板伪扫描与合成伪缺陷之上的两阶段'
    '仅正常样本（normal-only）流程确定。具体而言：利用完整管线使每个模板由其余模板打分，从而为每个'
    '模板生成一个无缺陷伪扫描；伪缺陷按⟦PO3AD⟧的增广实践，在模板表面以法向引导的局部'
    '形变合成。候选配置 (K, β, α) 按两个判据的排名之和排序：(i) 对象级行为——区分每类'
    '无缺陷伪扫描与其伪缺陷变体的类内 AUROC，按类别宏平均；(ii) 点级行为——伪缺陷点'
    '排在正常点之上的敏感度。任何选择步骤都不读取测试标签。',

    '搜索网格为 K ∈ {5, 10, 25, 50} × β ∈ {0, 0.25, 0.5, 1, 2} × α ∈ {none, 0.5, 1, '
    '2}。报告设定 (K, β, α) = (10, 1.0, 1.0) 在该协议执行之前已由诊断性敏感性扫描冻结'
    '；协议在三个种子上各自的选择、以及所选设定在正式协议下的表现见第 2.4 节；'
    '第 2.6 节将同一程序应用于另外两个基准。',
]

# 中文期刊正文一般不排伪代码浮动体（模板亦无算法环境，见优化意见 42）：
# 原"算法 1"删除，评分流程改由本节文字与图 1 三级框图表达。
S26 = [
    '对单个测试扫描 S，完整评分阶段依次经过三级处理：首先按式 (3) 利用预构建的公差表'
    '（式 (1)、(2)，池化规模 K）对各通道残差进行标准化，并按式 (4) 组合为逐模板证据；'
    '随后按式 (5)、(6) 以温度 β 完成软最小模板融合；最后按式 (9)、(10) 在测试扫描的 '
    'kNN 图上以强度 α 完成空间一致性正则化，得到逐点分数 s̃(x)。整个流程与图 1 中的'
    '三级框图一一对应，全程仅复用管线已有的残差矩阵、最近模板索引与邻居表。',
    '需要全分辨率分数时，通过存储的采样索引映射将 s̃ 传递回原始坐标。对象'
    '分数沿用基线的标准化 top-k 均值（k = 100）：',
    '@EQ11',
    '式 (11) 与基线保持一致，使全部改进可归因于第 1–3 级。退化输入按确定性规则处理：若扫描的'
    '分数场方差为零或有效分数点数少于 k，则记 S_{obj} = 0。每一级均为 O(N·K) 或 '
    'O(N·M) 的稠密张量运算，且仅作用于管线已计算好的量，因此 GCAS 在已存储的邻居表'
    '之外不新增内存占用，整体开销仍由配准主导；实测分解见第 2.5 节。',
]

S31 = [
    'Real3D-AD⟦Reg3D-AD⟧包含 12 个类别，每类提供 4 个无缺陷模板，'
    '共有 1 206 个测试扫描点云；如引言所述，全部点云均为高精度工业扫描仪对'
    '实物工件的实测数据，而非仿真生成。本文在对象级与点级分别报告受试者工作特征'
    '曲线下面积（Area Under the ROC Curve，AUROC）与平均精度（Average Precision，'
    'AP）；在该'
    '基准 1–5% 的低异常占比下，AP 是信息量更充分的点级汇总指标⟦DavisGoadrich⟧。点级'
    '指标遵循官方加载器坐标网格，模型输出分数通过确定性最近邻指派映射至真值坐标（细节'
    '见补充材料 E）。两处诊断性分析使用 PCD 原生 10K 点网格并已在文中注明（第 2.4、'
    '2.5 节）；两种网格的数值均包含在公开的逐样本输出中。',

    '在评测协议上，本文报告的所有数值均为整条管线（含配准）在种子 '
    '{42, 104771, 209500} 下'
    '三次独立重复实验的均值；逐种子数值与标准差见补充材料。所有随机流均由运行种子派生'
    '。超参数 (K, β, α) 全基准共用一套设定、从不按类别调参；第 1.5 节 normal-only '
    '协议的实际选择及其对结果的影响见第 2.4 节。',

    '首要对照基线是带有常规评分阶段的同一管线（K = 50 池化、fitness^{2} 融合、'
    '无空间正则化），以便准确隔离出 GCAS 的独立贡献。该基线的对象级 AUROC 达到 '
    '0.893，高于表 4 中的所有外部方法（最强为 DLF-3AD 预印本的 0.859，同行评审方法中'
    '最强为 Template3D-AD 的 0.844）。其每个组成部分均沿用文献成熟做法：基于 FPFH '
    '特征对应的 RANSAC 与 ICP 精化沿用 Reg3D-AD⟦Reg3D-AD⟧的配准流程；利用局部池化的'
    '跨模板统计量对残差进行标准化是 2D 检测中逐位置正常建模⟦PaDiM⟧向点云的自然推广；'
    'fitness 平方加权则对应第 1.3 节所分析的默认配准置信度加权方案；点级指标在补充'
    '材料 F 所列各网格中最严格的一种上评测，对象级协议为基准自带。本文另转录了 '
    '2023–2026 年间 19 个已发表方法的结果，逐方法协议注记见补充材料 F。',

    '在实现上，采用 Open3D⟦Open3D⟧ RANSAC（FPFH 对应，每模板 8 次重启）与点到点 ICP；'
    'N = 10 000 点；预处理、校准与正则化共享 k_{nn} = 50 邻居表；评分全程采用 float64 '
    '双精度浮点计算。全部计时均在单个 CPU 工作进程上实测，报告于第 2.5 节。',
]

S32 = [
    '表 1 报告正式协议（三种子、官方坐标网格、单一设定）下的全基准对照。',
    '@TAB1',
    '四项指标全面提升，增益最大处在点级：AP +0.101、AUROC +0.061。各项指标在不同随机'
    '种子间的波动均低于 0.006，比改进量小一个数量级；对 12 个类别的配对逐类均值做双侧 '
    'Wilcoxon 符号秩检验，四项指标均拒绝无差异的原假设（p = 0.021、0.034、0.0005、'
    '0.0024；补充材料 B）。本文两种方法的逐类结果见表 2 与表 3；逐种子数值与其余两项'
    '指标见补充材料 A 与 B。',

    '增益在各类别上表现一致。点级 AUROC 在全部 12 类上均实现提升，点级 AP 在 12 类中'
    '的 10 类提升；两个点级 AP 例外（diamond 0.697 → 0.681、fish 0.652 → 0.642）恰是'
    '常规阶段点级表现最强的类别，且两者的点级 AUROC 仍在提升（0.957 → 0.988 与 '
    '0.940 → 0.991）。对象级 AUROC 在 12 类中的 9 类提升，fish 持平，diamond 与 shell '
    '各下降 0.011。最大增益集中在常规阶段表现最弱的类别：airplane（对象 AUROC +0.096、'
    '点 AP +0.208）、starfish（+0.094、+0.256）、seahorse（+0.035、+0.193）——这些'
    '类别包含薄壁结构或小缺陷区域，正是粗粒度公差与均匀融合负面影响最显著之处。所有'
    '类别在各项指标上的性能降幅均未超过 0.028（单项最大降幅为 chicken 的对象 AP），'
    '而点级 AP 增益超过 0.05 的类别多达 8 个。',
]

S33 = [
    '已发表结果取自各自文献原文、在各自评测协议下测得，逐方法协议编目于补充材料 F；'
    '破折号（—）表示该指标未见报告。本文两种方法（常规管线与 GCAS）沿用第 2.2 节的'
    '三种子实验结果，在官方加载器坐标网格上测得——这是所涉各点级网格中最严格的一种。',
    '@TABLES234',
    '由表 4 可以得出以下三点观察。其一，在所有存在已发表结果的指标列中，GCAS 的宏均值'
    '均取得最高表现：现有最强外部方法为对象级 AUROC 0.859（DLF-3AD，2026 预印本；同行'
    '评审方法中最强为 Template3D-AD 的 0.844）与点级 AUROC 0.932（M2P-AD，2026 年同期'
    '预印本），前者低于常规基线的对象级 0.893，后者低于 GCAS 的点级 0.958。其二，文献'
    '中最强的点级结果（M2P-AD 0.932、Template3D-AD 0.925、Simple3D 0.923）各自在自有'
    '评估网格上测得，因此已发表方法之间的点级指标在严格意义上不可直接互比，而在这些最'
    '强条目中本文方法是唯一在全分辨率官方网格上评测的。其三，点级 AP 作为在 Real3D-AD '
    '低异常占比（占点数 1–5%）下对假阳性最敏感的指标，自 PointCore 之后的已发表方法均'
    '未报告（包含表 4 中全部 2025–2026 年工作）；本文评测协议在三种子下逐类报告全部四'
    '项指标。',
]

S34 = [
    '除特别说明外，本小节的消融实验采用表 1 的正式协议，各基准对照行严格对齐表 1。'
    '唯一例外是下文的平滑强度权衡分析：其在 airplane 与 chicken 上以单种子 10K 点 PCD 原生'
    '网格测试，其参考价值主要体现在各配置的相对排序而非绝对数值。',

    '首先考察第 2 级的融合算子：表 5 固定常规 K = 50 校准，不加空间正则化，只变动融合'
    '算子。',
    '@TAB5',
    '除朴素中位数外，所有鲁棒算子在对象级均优于 fitness^{2}，印证了常规加权方案的实际'
    '收益甚微。丢弃跨模板一致性的算子在点级性能上付出代价，这与第 1.3 节的理论分析相'
    '符：硬最小值使对象 AUROC 提升 0.011，但点 AUROC 下降（0.896 → 0.892）；中位数的'
    '两项点级指标均出现下滑。β ≤ 0.5 的软最小是唯一在四项指标上均实现改善的算子族'
    '（图 2(c)），其中 β = 0.25 在对象 AUROC、点 AUROC 与点 AP 上均为族内事后最优。在'
    '报告设定 β = 1 下点级 AUROC 较常规值低 0.002，该回退由空间正则化补回并反超'
    '（表 6）。报告设定取 β = 1；第 1.5 节已执行的 normal-only 协议在三个种子中的两个'
    '独立选中 β = 0.25，与该事后最优一致，且两种设定在完整三级栈下各项指标的差距均不'
    '超过 0.005（见下文）。',
    '图 3 进一步在全部 3×1 206 个正式运行扫描上验证了局限（2）'
    '的前提：归一化 fitness^{2} 权重接近均匀，这正是各逐点鲁棒替代方案均胜过它的原因。',
    # v13：图 2 流位自此后移两段（原在 @FIG3 之前）。根因：表 2/3/4 单栏
    # 夹层的表 4 按行跨页续排、表尾落在下一页页顶，而图 2 页顶通栏浮排的
    # 锚段同在该页时两者同占页顶——内嵌表行不避让浮动块（v10/v11 已证），
    # 图 2 直接叠印在表 4 末 11 行与表注上（云端 QA 第 8 页 99 处叠印）。
    # 后移至“其次考察池化规模 K”段（图 2(a) 的引用段）之后，锚段落入
    # 下一页，图 2 顶浮于表尾页的次页页顶；首引图 2(c) 在前页，先文后图不变。
    '其次考察第 1 级的池化规模 K。在完整三级架构下测试不同的池化规模：对象级 AUROC 宏均值依次'
    '为 0.924（K = 5）、0.928（K = 10）、0.924（K = 25）、0.919（K = 50），点 AP 依次'
    '为 0.562、0.580、0.585、0.589（图 2(a)）。对象级指标在 K = 10 处达到峰值，K = 5 '
    '因公差统计噪声过大而出现退化，点 AP 则随 K 稳步上升；K = 10 为对象级最优，其点 '
    'AP 与最大池化设定的差距不超过 0.010。实验表明多尺度融合的效果弱于单一最优尺度'
    '（补充材料 D）。',
    '@FIG2',
    # v13：表 6 流位定于此（v12 在“对于第 3 级”段后）。试排两处均叠印：
    # 留在“对于第 3 级”段后落进图 3 栏底浮矩形（v2 第 9 页 22 处），前移到
    # 表尾页栏底则跨栏断行、尾行漂进图 2 通栏带下（v2c 第 9 页 1 处）。置于
    # 图 2 锚段后，表 6 整块落在显示页左栏中部——上距图 2 带下缘、下距
    # 图 3 顶边均有正文缓冲。首次正文引用（表 6）在前页，先文后表不变。
    '@TAB6',

    '对于第 3 级的平滑强度 α，在 airplane 与 chicken 上将其由 0.5 增至 2.0 时，点级 '
    'AUROC 单调上升而对象级 AUROC 下降，迭代扩散延续了这一趋势（图 2(b)；数值见补充'
    '材料 D）。这表明增强平滑能够持续改善点级排序，但会压平对象级 top-k 统计量所依赖'
    '的分布上尾；α = 1 在两项指标间取得平衡。',
    # v13：图 3 锚位自“图 3 进一步”段后（表尾页末段）移至此。原锚段是
    # 表尾页最后一段而显示页在次页，触发 LibreOffice 类“浮动屏障”（锚页
    # 右栏 4.7 cm 悬空，v2 渲染实测）；移至本段后锚段与显示页同页（v9
    # 原则），图 3 仍栏底浮排。首引“图 3”在前页，先文后图不变。
    '@FIG3',

    '为考察各级的独立贡献，表 6 在正式协议下沿主路径逐级累加各模块；完整 2×2×2 开关'
    '网格见补充'
    '材料 D。软最小融合贡献了消融路径上首步的对象级增益（较常规阶段 +0.012 AUROC、+0.014 '
    'AP）；细粒度 K = 10 进一步提升 +0.009 对象 AUROC；空间正则化随后贡献了点级增益的'
    '主要部分（点 AUROC +0.065、点 AP +0.095），并将对象 AUROC 进一步提高 +0.014。'
    '任何单独一级均无法完全解释表 1 的总增益：对象级增益在三个新增组件上逐级累积，而'
    '点级性能的大幅跃升则主要由最后一级提供（图 4）。',
    '此外，在三个随机种子上分别运行第 1.5 节的 normal-only 协议以检验参数选择：两个'
    '种子选出 (K, β, α) = (25, 0.25, 1.0)，第三个选出 (25, 1.0, 2.0)。在正式协议下，'
    '多数票所选设定的四项指标达到 0.923/0.932/0.958/0.582，与报告设定的 0.928/0.935/'
    '0.958/0.580 极为接近；第三个种子的选择表现相当（0.924/0.932/0.960/0.585）。按'
    '报告精度各指标差距均不超过 0.005，且两种设定均完整保留了表 1 的全部增益，表明'
    '实验结论并不敏感依赖于参数究竟源自无标签选择协议还是诊断扫描所冻结的设定。残余的'
    ' K 分歧（25 对 10）与伪缺陷合成的尺寸偏置相符——其形变区域宽于最具挑战性的真实'
    '缺陷（补充材料 D）。',
    '最后考察模板数 M 的影响：在全部 C(4, M) 个模板子集上重建公差表并重新执行融合（图 5；'
    '逐 M 数值见补充'
    '材料 D）。GCAS 的点级增益在任意 M 下均成立；对象级优势在 M = 1 与 M ≥ 3 时成立，'
    '仅在 M = 2 时未能显现——当仅有两个模板时，软最小融合缺乏足够的交叉证据来区分'
    '配准失败与真实缺陷，因为第 1.3 节的共识机制依赖于多模板的多数表决效应。',
    # v13：图 4/图 5 图带流位自“此外…”段后再移至节尾（图 5 的引用段之后）。
    # 图 2 错页引发回流后，原锚段收尾于图带显示页的前一页，屏障空洞移相
    # 复现（v2 渲染实测锚页右栏 4.57 cm 悬空）；移至“最后考察模板数”段
    # 后，锚段与显示页同页（v9 原则）。图 4 首引在“为考察各级”段、图 5
    # 首引在本段，先文后图不变。
    '@FIGBAND45',
]

S35 = [
    '在单个 CPU 核心上实测（种子 42，airplane，100 个扫描），完整 GCAS 评分阶段耗时 '
    '21.4 ms/样本，常规阶段为 3.0 ms/样本。端到端推理仍由配准主导：在全部 3 618 个'
    '正式运行样本上汇总，单 CPU 工作进程的实际运行时间为 15.9 s/样本（墙钟时间），'
    '其中仅配准环节即占 8.5 s（完整分解见补充材料 E）。因此评分阶段引入的额外耗时不足'
    '总推理时间的 1%。',

    '图 6 展示定性点分数图（分数由低到高按浅灰至红色着色；每个面板中的蓝色虚线框标出 '
    '真值缺陷区域）。在薄壁结构类别中，常规评分阶段沿边缘与配准错位带容易产生弥散的'
    '虚警响应，而 GCAS 能够有效抑制此类假阳性，并增强真实缺陷区域的响应一致性。',
    '图 7 记录三个最弱类别上的主导失败模式（着色约定与 GT 虚线框同图 6）。对 gemstone'
    '、duck、shell，各取 GCAS 下逐样本点级 AP 最低的异常样本（种子 42，10K 点网格）。'
    '三例均为仅占表面 0.5–1.1% 的浅凹陷（sink）缺陷，校准后缺陷区域仅被轻微拉升：其组合残差'
    '中位数位于背景分布的第 61 至 94 百分位，落在正常公差带内而非上尾区；最高分反而被'
    '标注框外的伪响应占据（gemstone 与 duck 为稀疏扫描边界带，shell 为内部残差纹理），'
    '导致查准率在 AP 指标所奖励的工作点处急剧下降。残差通道本身几乎无法区分这些浅 '
    'sink 缺陷与正常公差，此类失败案例明确界定了仅依赖评分阶段所能改善的性能边界；其'
    '进一步突破有赖于引入更为丰富的残差特征通道（第 3 节）。',
]

S36 = [
    '本文通过另外两个基准检验 GCAS 所针对的三处局限是否为 Real3D-AD 所独有。'
    'Anomaly-ShapeNet⟦IMRNet⟧与 MulSen-AD⟦MulSenAD⟧均提供完整 3D 扫描；深度扫描基准 '
    'MVTec 3D-AD 与 Eyecandies⟦MVTec3D,Eyecandies⟧为单视角深度图像、且每类包含数百个'
    '正常训练样本，不属于 GCAS 所针对的少模板完整 3D 配准场景，故不在讨论范围。在这两'
    '个基准上，本文首先将冻结管线与默认参数 (10, 1, 1) 直接迁移测试，随后仅利用目标'
    '基准提供的正常模板，重复第 1.5 节的无标签选参流程（表 7）。',
    '在 Anomaly-ShapeNet 上，固定设定使点 AUROC'
    ' 与点 AP 分别在 39/40 与 35/40 类上提升，宏均值达到 0.955 与 0.608；对象 AUROC '
    '则从 0.904 降至 0.832，反转主要集中于结构对称、配准质量较差的类别（cup0 '
    '1.00 → 0.51、headset1 0.93 → 0.25）。这两种效应均符合第 1.2 节与第 2.4 节的机制'
    '分析：该基准中的缺陷多为大面积凸起（bulge）与凹陷（concavity），在此尺度下 '
    'K = 10 无法体现相较于粗池化的优势，反而会放大公差表的统计噪声；而过强的平滑操作'
    '又容易将正常的残差纹理连成片状虚警，进而被 top-k 统计量误判。此外该数据集上的'
    '配准难度更高（fitness 均值 0.49，对照 0.59）。可见在缺陷尺度存在显著差异的数据集'
    '上，单一的全基准统一设定并非对象级最优。',

    '进一步地，仅用 Anomaly-ShapeNet 的模板、按随机种子独立运行第 2.5 '
    '节协议，各次运行均自适应选出了更粗粒度的池化规模——三个种子全部选中 K = 50 与 '
    'α = 1，其中两个选中 β = 2——这与基于缺陷尺度的机制分析完全吻合。按各种子自身的'
    '选择重新评测，对象级性能恢复并反超常规阶段（0.832 → 0.926 ± 0.006，对照 0.904），'
    '点级领先优势同步扩大，最严重的性能反转在每个种子上均得到修复（cup0 0.51 → 1.00、'
    'headset1 0.25 → 0.87；图 8）。与该基准上按类别训练的方法相比（补充材料 G），重选'
    '参数后的对象级性能与 2025–2026 年间的最优方法相当（SeDiR 0.933、DLF-3AD 0.921、'
    'PASDF 0.900），而 GCAS 在两种设定下的点级 AUROC 均超过了该基准上所有已发表的点级'
    '结果（最强为 M2P-AD 的 0.943，自有网格）。固定设定行则给出了不进行参数重选时的'
    '性能下界。',
    # v13：表 7 流位自“重选”段前移至其后。v13 回流后表 7 尾行落进图 4/图 5
    # 底浮图带右格（v2b 渲染实测第 10 页 14 处叠印）且跨页断开（表注孤悬
    # 次页页顶）；后移一段后整表进入次页栏顶、在图 6 通栏矩形之上整块容纳。
    # 首次正文引用（表 7）在本节首段，先文后表不变。
    '@TAB7',
    'MulSen-AD 则揭示了另一类适用边界。该基准提供由高精度激光扫描仪等多种传感器对'
    '实物采集的完整 3D 扫描⟦MulSenAD⟧，但每类拥有上百个正常'
    '训练样本'
    '而非少量模板；本文选取每类训练扫描中按文件编号升序排列的前 4 个样本作为模板（在接触'
    '任何标注前即已固定），并将冻结管线直接迁移。在其 644 个测试扫描中，152 个异常'
    '扫描因缺陷在几何上不可见（如纯颜色缺陷）而缺少点级真值，这些样本仅进入对象级'
    '评测。点级性能在两种设定下再次成功迁移：固定设定下分别在 14/15 与 13/15 类上'
    '提升，重选后为 15/15 与 14/15，且重选再次选出了更粗的池化规模。但对象级指标在'
    '此次仅实现部分恢复（0.648，对照 0.677），未能像在 Anomaly-ShapeNet 上那样实现'
    '反超。逐样本分析表明，剩余差距主要集中于违背管线刚性假设、而非参数失配的类别，'
    '其中 cotton 最为典型：织物每次扫描的披挂形态各不相同，每个正常扫描均包含样本特有'
    '的褶皱而同时偏离全部四个模板，软最小共识在所探查的任何 (K, β, α) 下都会将这种'
    '偏离判定为异常（补充材料 H）。这属于评分阶段上游所依赖的刚性配准假设失效，无法仅'
    '通过调整评分参数予以修复。',
    # txxb-v1 QA（七轮定稿）：图 6（通栏顶浮，显示第 12 页页顶）与图 7/图 8
    # 通栏底浮图带（两栏宽底浮块锚位相邻会叠印——LibreOffice 类引擎不执行
    # 浮动表互斥（v10 教训），按 v9 图 4/图 5 先例合带、格内垂直居中，顶底
    # 分层互不争位）锚位置于 MulSen 首段之后：表 7 前移后重选段结尾在第 11
    # 页右栏 70% 处，MulSen 首段恰好填满该栏并自然溢入第 12 页，锚段随之
    # 落在显示页内（v9 原则）。三图首次引用均在前页，先文后图不变。
    '@FIG6',

    '由此可以明确纯几何检测的两条边界。其一，几何不可见缺陷——在此处占异常扫描的 '
    '30.8%（494 个异常扫描中的 152 个），且该基准中仅 4.3% 的缺陷只能由点云传感器检出'
    '⟦MulSenAD⟧——构成了所有纯几何点云方法对象级性能的理论上限；其二，可形变物体从'
    '根本上违背了刚性配准前提。与该基准自带的纯点云基线（均使用全部正常训练集而非仅 '
    '4 个模板）相比，本文常规阶段处于中游水平，而 GCAS 在两种设定下的点级指标均超越了'
    '该基准所报告的点云定位均值（补充材料 H）。',
]

S4 = [
    '本文重新审视了基于配准的 3D 异常检测中的评分阶段，指出现有常规设计存在的三处'
    '局限：对小缺陷而言池化过粗的公差统计、无法察觉逐点配准失败的逐模板 fitness 权重、'
    '以及与缺陷连续性相悖的逐点独立打分。GCAS 通过轻量级免训练算子逐一解决了上述问题'
    '——包括细粒度公差校准、软最小模板融合与 kNN 图空间一致性正则化，同时完全保持'
    '原有的配准、残差提取与聚合环节不变。在 Real3D-AD 三种子官方网格协议与全基准单一'
    '设定下，相较于同一条常规管线，GCAS 将对象级 AUROC 从 0.893 提升至 0.928、点级 '
    'AP 从 0.479 提升至 0.580，且每一级的贡献均经过了单独验证。无标签重选流程将该模块'
    '成功推广至 Anomaly-ShapeNet，而 MulSen-AD 实验则进一步界定了其适用边界'
    '（第 2.6 节）。这表明仅通过优化评分阶段，即可在零训练成本下获得显著的性能提升。',

    '在局限方面，本文默认参数由诊断性扫描一次性冻结、全基准统一使用，虽具通用性，'
    '但尚未做到完全无标签自适应；此外，GCAS 的整体处理速度受限于上游配准环节，'
    '单样本耗时约 16 s。评分阶段仅能重组与挖掘残差通道所提供的信息，这正是第 2.5 节'
    '浅凹陷失败案例'
    '与第 2.6 节几何不可见缺陷共同划定的性能边界。在后续工作中引入更为丰富的几何与'
    '语义残差通道（如预训练基础模型的冻结特征），将是有价值的研究方向。',
]

CONTRIB = [
    '【作者一姓名】：方法的提出，论文构思和撰写；',
    '【作者二姓名】：实验的设计及数据整理和分析；',
    '【作者三姓名】：论文审核与编辑写作。',
]

SUPP_NOTE = (
    '本文附录 A–I（逐种子结果、逐类明细表、公差构建自对敏感性、组件消融与五项负对照、'
    '可复现清单与端到端计时分解、逐方法对比协议编目、Anomaly-ShapeNet 40 类逐类对比、'
    'MulSen-AD 细节、命题 1 证明）及三张附图（逐级定性分数图、效率分解图、几何可见性'
    '分解图）拟作为补充材料随投稿一并提供，正文仅保留其结论性引用；如版面受限，'
    '该部分在投稿版中可整体移出。'
)

# 约 300 字（模板要求）；开源承诺移至"数据与程序可用性"段
INNOVATION = (
    '本文的创新点在于：（1）首次系统指出基于配准的三维异常检测中常规评分阶段的三处'
    '结构性局限：公差在远大于小缺陷的邻域上池化、全局适配度加权无法抑制局部错配模板、'
    '逐点独立打分忽略缺陷的空间连续性；（2）提出免训练评分模块 GCAS，以细粒度公差'
    '校准、软最小模板融合与 kNN 图空间一致性正则化三级算子逐一解决上述局限，软最小'
    '融合具有可证明的共识保真、插值与单模板污染上界性质，新增耗时不足端到端推理的 '
    '1%；（3）在 Real3D-AD 上以三种子、官方坐标网格与全基准单一设定取得对象级 AUROC '
    '0.928、点级 AUROC 0.958，均超过已发表最优结果，点级增益可迁移至另两个公开基准；'
    '（4）提出仅用正常模板的无标签参数选择协议，支持在新基准上自适应重选参数。'
)

DATA_AVAIL = (
    '本文全部脚本、冻结参数、逐样本输出与图表生成代码公开于：【仓库 URL 待作者补充'
    '（投稿时应已可访问，匿名或实名均可，与编辑部确认）】，许可证【待定】。'
    'Real3D-AD、Anomaly-ShapeNet 与 MulSen-AD 均为公开基准数据集，获取方式见各自原文。'
)

# 一式一号（模板要求 6）：原 v3 的 8 个编号中 5 个为复合式，v4 拆分为 11 式；
# 原式 (3) 的通道权重取值与原式 (8) 的 k = 100 移入正文行文。
EQS = {
    1: 'μ_{t}^{(c)}(p) = mean A_{t}^{(c)}(p)，',
    2: 'σ_{t}^{(c)}(p) = std A_{t}^{(c)}(p)，',
    3: 'z_{t}^{(c)}(x) = ( r_{t}^{(c)}(x) − μ_{t}^{(c)}(q_{t}(x)) ) / '
       '( σ_{t}^{(c)}(q_{t}(x)) + ε )，',
    4: 'z_{t}(x) = Σ_{c} λ_{c}·z_{t}^{(c)}(x)，',
    5: 'w_{t}(x) = exp(−β z_{t}(x)) / Σ_{j} exp(−β z_{j}(x))，',
    6: 's(x) = Σ_{t} w_{t}(x) z_{t}(x)，',
    7: 'β → 0 时，s(x) → (1/M) Σ_{t} z_{t}(x)，',
    8: 'β → ∞ 时，s(x) → min_{t} z_{t}(x)，',
    9: 'σ_{S} = α·median{ d(x, y) : x ∈ S, y ∈ N(x) \\ {x} }，',
    10: 's̃(x) = Σ_{y∈N(x)} exp(−d(x,y)^{2}/2σ_{S}^{2})·s(y) / '
        'Σ_{y∈N(x)} exp(−d(x,y)^{2}/2σ_{S}^{2})。',
    11: 'S_{obj} = ( mean top-k s̃ − mean s̃ ) / std s̃，',
}

# 编号: (png, 宽度, 中文题, 中文分图/图例说明, 英文题, 英文分图/图例说明)。
# 题注体例对照 OEP 模板样例与已刊文章（如"图4 CA 模块示意图/Fig. 4 Schematic
# diagram of CA module"）：简短名词短语、不带冒号从句；分图说明中英对照成行。
FIGS = {
    1: ('fig1_pipeline.png', FULL_W,
        '图1 基于配准骨架的 GCAS 管线总览', '',
        'Fig. 1 Overview of the registration-based pipeline with GCAS', ''),
    2: ('fig2_sensitivity.png', FULL_W,
        '图2 参数敏感性分析',
        '（a）池化规模 K 扫描；（b）平滑强度权衡；（c）融合算子族',
        'Fig. 2 Parameter sensitivity analysis',
        '(a) Pooling size K sweep; (b) smoothing strength trade-off; '
        '(c) fusion operator family'),
    3: ('fig4_weights.png', COL_W,
        '图3 最大归一化 fitness^{2} 模板权重的分布', '',
        'Fig. 3 Distribution of the largest normalized fitness^{2} template '
        'weight', ''),
    4: ('fig6_component.png', COL_W,
        '图4 主路径逐级累积消融',
        '（柱：对象级 AUROC 逐级累积；折线：点级 AP）',
        'Fig. 4 Per-stage cumulative ablation along the main path',
        '(Bars: object-level AUROC; line: point-level AP)'),
    5: ('fig7_templates.png', COL_W,
        '图5 全部 C(4, M) 模板子集上的模板数曲线',
        '（a）对象级 AUROC；（b）点级 AUROC',
        'Fig. 5 Template-count curves over all C(4, M) template subsets',
        '(a) Object-level AUROC; (b) point-level AUROC'),
    6: ('fig3_qualitative.png', FULL_W,
        '图6 五个类别上的定性点分数图',
        '（上：真值；中：常规评分阶段；下：GCAS；蓝色虚线框为真值缺陷区域）',
        'Fig. 6 Qualitative point-score maps on five categories',
        '(Top: ground truth; middle: conventional stage; bottom: GCAS; '
        'blue dashed boxes mark ground-truth defect regions)'),
    7: ('fig5_failure.png', COL_W,
        '图7 三个最弱类别上的失败案例',
        '（浅凹陷（sink）缺陷；着色与虚线框约定同图 6）',
        'Fig. 7 Failure cases on the three weakest categories',
        '(Shallow sink defects; colormap and dashed boxes as in Fig. 6)'),
    8: ('fig8_asn.png', COL_W,
        '图8 向 Anomaly-ShapeNet 的迁移结果',
        '（a）对象级 AUROC；（b）点级 AP',
        'Fig. 8 Transfer results on Anomaly-ShapeNet',
        '(a) Object-level AUROC; (b) point-level AP'),
}

TABLE_META = {
    'T1': ('表1 12 类宏均值（三种子均值±标准差，官方坐标网格）',
           'Table 1 Macro average over 12 categories (mean±std over three '
           'seeds, official coordinate grid)',
           'Conventional = 带常规评分阶段的同一管线。Δ = GCAS − Conventional，由未'
           '舍入三种子均值计算；std 为种子间总体标准差。', COL_W_TW, 6.5, 0.28),
    'T2': ('表2 逐类对象级 AUROC（Real3D-AD）',
           'Table 2 Object-level AUROC by category (Real3D-AD)',
           '3DKeyAD 展示其最优已发表变体（Raw+FPFH+ISS+FS）；PointCore 取全指标配置'
           '（PointMAE+Raw，其表 II）。本文两行为表 1 的三种子均值；逐方法协议见'
           '补充材料 F。', FULL_W_TW, 7.5, 0.16),
    'T3': ('表3 逐类点级 AUROC（Real3D-AD）',
           'Table 3 Point-level AUROC by category (Real3D-AD)',
           '各来源的网格定义不同，记录于补充材料 F；本文两种方法均使用官方加载器坐标'
           '网格。无逐类点级结果的方法此处省略、见表 4：其中 IMRNet 与 PO3AD 在 '
           'Real3D-AD 上未报告点级结果，R3D-AD 与 Template3D-AD 仅报告宏观点级均值，'
           'SeDiR 的逐类数值仅见于其补充材料（未转录）。', FULL_W_TW, 7.5, 0.16),
    'T4': ('表4 12 类宏汇总',
           'Table 4 Macro summary over 12 categories',
           '破折号（—）表示来源未报告该指标。R3D-AD 的宏 P-AUROC（‡）来自其 '
           'Real3D-AD 消融实验（其表 3 完整模型 D），是该文唯一报告处。已发表方法中'
           '只有 Reg3D-AD、M3DM（均经 Real3D-AD 基准套件）、Group3AD 与 PointCore '
           '报告过 AP，本文常规基线在官方网格下的点级 AP 0.479 接近文献最高值'
           '（0.251）的两倍。', FULL_W_TW, 7.5, 0.18),
    'T5': ('表5 融合算子比较（三种子均值±标准差，官方坐标网格）',
           'Table 5 Fusion operator comparison (three-seed mean±std, official '
           'coordinate grid)',
           'K = 50 校准、无空间正则化。', COL_W_TW, 6.5, 0.30),
    'T6': ('表6 主路径逐级累积消融（三种子均值±标准差，官方坐标网格）',
           'Table 6 Per-stage cumulative ablation along the main path '
           '(three-seed mean±std, official coordinate grid)',
           '第 1 行是表 1 的常规阶段，第 4 行是完整 GCAS。', COL_W_TW, 6.5, 0.34),
    'T7': ('表7 官方坐标网格下的迁移表现',
           'Table 7 Transfer performance on the official coordinate grids',
           '“固定”指把冻结的 Real3D-AD 设定 (10, 1, 1) 直接沿用；“重选”指只用目标基准'
           '的模板运行第 1.5 节的无标签协议，不触及任何测试数据。Anomaly-ShapeNet 为'
           '三种子均值（各格标准差至多 0.006，唯重选设定的点 AP 为 0.012）；'
           'MulSen-AD 为单提取种子。加粗为各基准区块内该列最优值。',
           COL_W_TW, 6.5, 0.34),
}


# ---------------------------------------------------------------- 构建

def emit_table(doc, tag, tables):
    if tag in FLOAT_TABLES:      # v10：栏宽表 5/6/7 浮动容器化（见文件头注记）
        add_floating_table(doc, tag, tables)
        return
    zh, en, note, tw, size, ratio = TABLE_META[tag]
    caption_pair(doc, zh, en, size=9)
    rows = tables[tag]
    group_rows = tuple(i for i, r in enumerate(rows)
                       if r[0].startswith('*') and all(not c for c in r[1:]))
    add_3line_table(doc, rows, tw, size, first_col_ratio=ratio,
                    group_rows=group_rows)
    table_note(doc, note)


def emit_stream(doc, paras, tables, indent=True):
    for item in paras:
        if item.startswith('@EQ'):
            n = int(item[3:])
            equation(doc, EQS[n], n)
        elif item.startswith('@TAB') and item != '@TABLES234':
            emit_table(doc, 'T' + item[4:], tables)
        elif item == '@FIGBAND78':
            # txxb-v1 QA（八轮）：图带（6.4 cm）顶浮 + 图 6（10.5 cm）底浮
            # 同页顶底换槽——两块合高 17 cm 恰容于一页，正文夹排其间
            add_floating_figure_band(doc, FIGS[7], FIGS[8], y_spec='top')
        elif item == '@FIG6':
            add_floating_figure(doc, *FIGS[6], y_override='bottom')
        elif item.startswith('@FIG'):
            # 按图宽自动选浮排方式：通栏图页顶浮排，栏宽图栏底浮排（v8/v9）
            fig = FIGS[int(item[4:])]
            add_floating_figure(doc, *fig, column=(fig[1] != FULL_W))
        else:
            body_para(doc, item, indent=indent)


FANG = '仿宋'


def build(blind=False):
    """blind=True 生成盲审稿：删除中英文作者姓名行、单位行与首页脚注中的
    作者简介/通信作者内容（投稿指南 2(1)），其余与原稿逐字一致。"""
    entries = parse_bib(BIB)
    tables = load_md_tables(DRAFT)

    doc = Document(TEMPLATE)

    # —— 模板 Normal 样式首行缩进清零（云端 QA 定位的表格断行真根因）——
    # 图学学报模板（LibreOffice 自 .doc 转换）的 Normal 样式自带
    # w:ind firstLine=420：表格单元格/公式/题注等未显式设缩进的段落全部
    # 继承 2 字符首行缩进，622 twip 数据列刨掉 420 只剩 0.33 cm，数值
    # "0.716" 在 "0." 后断行、表头 airplane 断成 air/plane。正文首行缩进
    # 由 body_para 的 fli_chars=2 逐段显式给出，样式级清零不影响正文。
    # 验证：_qa_cloud/render_patchtest/pt2.pdf（清零后表 2/3/4 全部单行）。
    for style in doc.styles.element.findall(qn('w:style')):
        if style.get(qn('w:styleId')) == 'Normal':
            ppr = style.find(qn('w:pPr'))
            if ppr is not None:
                for ind in ppr.findall(qn('w:ind')):
                    ppr.remove(ind)

    # —— v13：页眉边框净化（模板 .doc→.docx 转换残线，见函数 docstring）——
    _sanitize_headers(doc)

    # —— 收集模板分节符克隆 ——
    para_sects = [s for p in doc.element.body.iterchildren(qn('w:p'))
                  for s in p.findall(qn('w:pPr') + '/' + qn('w:sectPr'))]
    sect_front = None
    sect_2col = None
    sect_1col = None
    for s in para_sects:
        cols = s.find(qn('w:cols'))
        num = cols.get(qn('w:num')) if cols is not None else None
        if sect_front is None:
            sect_front = copy.deepcopy(s)
        elif num == '2' and sect_2col is None:
            sect_2col = copy.deepcopy(s)
        elif num != '2' and sect_1col is None:
            sect_1col = copy.deepcopy(s)
        if sect_front is not None and sect_2col is not None \
                and sect_1col is not None:
            break
    assert sect_front is not None and sect_2col is not None \
        and sect_1col is not None, '模板分节符提取失败'
    # 模板段落级双栏节为示例区的 0.9 cm 大栏距（space=1294）；正文栏距
    # 按模板 body 末级 sectPr 的 438 twip（0.77 cm）统一，与 COL_W_TW 一致
    cols = sect_2col.find(qn('w:cols'))
    cols.set(qn('w:space'), '438')

    # —— 清空正文（保留 body 末级 sectPr）——
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    # ============ 第 1 节：单栏头部（TXXB 模板版式） ============
    # 中文题名：二号（22 pt）黑体居中
    p = doc.add_paragraph()
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=6, after=6,
         line=360)
    emit_runs(p, ZH_TITLE, size=22, ea=HEI)

    # 盲审稿以等高空白段替代作者/单位行（内容删除、版面高度保留），使两
    # 版本正文分页逐页一致，浮动图锚位一套通用（投稿指南 2(1) 只要求删除
    # 内容；盲审版删除区留白为通行做法）
    p = doc.add_paragraph()
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=4,
         after=3, line=280)
    if not blind:
        emit_runs(p, '【作者一姓名】^{1}，【作者二姓名】^{1,2}，【作者三姓名】^{2}',
                  size=14, ea=FANG)
    else:
        r = p.add_run(' ')
        set_font(r, TNR, FANG, 14)
    p = doc.add_paragraph()
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, after=6,
         line=240)
    if not blind:
        emit_runs(p, '（1. 【第一单位至学院或系】，【所在省】 【所在市】 【邮编】；'
                  '2. 【第二单位】，【所在省】 【所在市】 【邮编】）',
                  size=9, ea=FANG)
    else:
        r = p.add_run(' ')
        set_font(r, TNR, FANG, 9)

    # 摘要：小 5 号（9 pt）楷体，两边各缩进 2 字符；标签黑体
    p = doc.add_paragraph()
    pset(p, snap=False, before=4, line=240, ind_chars_lr=2)
    emit_runs(p, '摘要：', size=9, ea=HEI, ascii_=HEI)
    r = p.add_run(ABSTRACT_ZH)
    set_font(r, TNR, KAI, 9)

    # 关键词：小 5 号楷体，两边各缩进 2 字符；≥5 个
    p = doc.add_paragraph()
    pset(p, snap=False, line=240, ind_chars_lr=2)
    emit_runs(p, '关键词：', size=9, ea=HEI, ascii_=HEI)
    r = p.add_run(KEYWORDS_ZH)
    set_font(r, TNR, KAI, 9)

    # 标签两行：中图分类号/DOI、文献标识码/文章编号（黑体标签 + TNR 值）
    p = doc.add_paragraph()
    pset(p, snap=False, before=3, line=240, ind_chars_lr=2)
    emit_runs(p, '中图分类号：', size=9, ea=HEI, ascii_=HEI)
    r = p.add_run('TP 391')
    set_font(r, TNR, SONG, 9)
    r = p.add_run('\u3000\u3000\u3000\u3000')
    set_font(r, TNR, SONG, 9)
    emit_runs(p, 'DOI：', size=9, ea=HEI, ascii_=HEI)
    r = p.add_run('10.11996/JG.j.2095-302X.0000000000')
    set_font(r, TNR, SONG, 9)

    p = doc.add_paragraph()
    pset(p, snap=False, line=240, after=8, ind_chars_lr=2)
    emit_runs(p, '文献标识码：', size=9, ea=HEI, ascii_=HEI)
    r = p.add_run('A')
    set_font(r, TNR, SONG, 9)
    r = p.add_run('\u3000\u3000\u3000\u3000\u3000\u3000')
    set_font(r, TNR, SONG, 9)
    emit_runs(p, '文 章 编 号：', size=9, ea=HEI, ascii_=HEI)
    r = p.add_run('2095-302X(0000)00-0000-00')
    set_font(r, TNR, SONG, 9)

    # 英文题名：4 号（14 pt）TNR 加粗居中，sentence case
    p = doc.add_paragraph()
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, before=10, after=6,
         line=300)
    emit_runs(p, EN_TITLE, size=14, ea=SONG, bold=True)

    p = doc.add_paragraph()
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, after=3,
         line=260)
    if not blind:
        emit_runs(p, '【SURNAME Given-name】^{1}, 【SURNAME Given-name】^{1,2}, '
                  '【SURNAME Given-name】^{2}', size=10.5, ea=SONG)
    else:
        r = p.add_run(' ')
        set_font(r, TNR, SONG, 10.5)
    p = doc.add_paragraph()
    pset(p, align=WD_ALIGN_PARAGRAPH.CENTER, snap=False, after=6,
         line=220)
    if not blind:
        emit_runs(p, '(1. 【Department, University】, 【City】 【Province】 '
                  '【Postcode】, China; 2. 【Department, University】, '
                  '【City】 【Province】 【Postcode】, China)',
                  size=7.5, ea=SONG)
    else:
        r = p.add_run(' ')
        set_font(r, TNR, SONG, 7.5)

    # Abstract：小 5 号（9 pt）TNR，两边各缩进 2 字符，≥2 500 字符
    p = doc.add_paragraph()
    pset(p, snap=False, line=240, ind_chars_lr=2)
    r = p.add_run('Abstract: ')
    set_font(r, TNR, HEI, 9, bold=True)
    r = p.add_run(ABSTRACT_EN)
    set_font(r, TNR, SONG, 9)

    # Keywords：与中文关键词完全对照
    p = doc.add_paragraph()
    pset(p, snap=False, line=240, after=6, ind_chars_lr=2)
    r = p.add_run('Keywords: ')
    set_font(r, TNR, HEI, 9, bold=True)
    r = p.add_run(KEYWORDS_EN)
    set_font(r, TNR, SONG, 9)

    add_sect_break(doc, sect_front)          # 结束单栏头部

    # ============ 双栏正文 ============
    # TXXB 已刊惯例（DOI …2026010029 与 …2026020380 实测）：引言不设标题、
    # 不编号，Keywords 后直接排引言正文；首个编号标题从「1」开始
    emit_stream(doc, INTRO, tables)

    heading1(doc, '1 几何一致异常评分（GCAS）')
    heading2(doc, '1.1 问题设定与管线总览')
    emit_stream(doc, S21, tables)

    heading2(doc, '1.2 第 1 级：细粒度公差校准')
    # txxb-v1 QA：图 1 锚位由 1.1 节末推迟至 1.2 节第 2 段之后。TXXB 头区
    # 与固定行距使分页较 OEP 前移，原锚位落第 3 页右栏而图显示第 4 页页顶，
    # 锚页与显示页错开触发"浮动屏障"（第 3 页右栏 12.4 cm + 第 4 页左栏
    # 17.0 cm 空白）；锚位与显示页同页后正文正常回填（v9 原则，与 JIG 版
    # 同一修复）。图 1 首次引用在 1.1 节（前页），先文后图不变。
    emit_stream(doc, S22[:6], tables)
    add_floating_figure(doc, *FIGS[1])
    emit_stream(doc, S22[6:], tables)
    heading2(doc, '1.3 第 2 级：软最小模板融合')
    emit_stream(doc, S23, tables)
    heading2(doc, '1.4 第 3 级：kNN 图空间一致性正则化')
    emit_stream(doc, S24, tables)
    heading2(doc, '1.5 参数选择协议')
    emit_stream(doc, S25, tables)
    heading2(doc, '1.6 评分流程、输出、复杂度与成本')
    emit_stream(doc, S26, tables)

    heading1(doc, '2 实验与结果分析')
    heading2(doc, '2.1 实验设置')
    emit_stream(doc, S31, tables)
    heading2(doc, '2.2 主结果')
    for item in S32:
        if item == '@TAB1':
            emit_table(doc, 'T1', tables)
        else:
            body_para(doc, item)
    heading2(doc, '2.3 与已发表方法的对比')
    for item in S33:
        if item == '@TABLES234':
            add_sect_break(doc, sect_2col)   # 通栏插排：表 2/3/4
            emit_table(doc, 'T2', tables)
            emit_table(doc, 'T3', tables)
            emit_table(doc, 'T4', tables)
            add_sect_break(doc, sect_1col)
        else:
            body_para(doc, item)
    heading2(doc, '2.4 消融实验')
    for item in S34:
        if item == '@FIG2':
            # v7：图 2 页顶通栏浮排（与图 6 同机制，见文件头注记）
            add_floating_figure(doc, *FIGS[2])
        elif item == '@FIGBAND45':
            # v9 合带：图 4/图 5 通栏底浮图带；v12 流位回到 3.4 节
            add_floating_figure_band(doc, FIGS[4], FIGS[5])
        elif item.startswith('@'):
            emit_stream(doc, [item], tables)
        else:
            body_para(doc, item)
    heading2(doc, '2.5 效率与定性分析')
    for item in S35:
        if item == '@FIG6':
            # v6：通栏浮动图文框（页顶浮排、跨双栏、无分节空白，见文件头注记）
            add_floating_figure(doc, *FIGS[6])
        elif item.startswith('@'):
            emit_stream(doc, [item], tables)
        else:
            body_para(doc, item)
    heading2(doc, '2.6 跨基准迁移')
    emit_stream(doc, S36, tables)

    heading1(doc, '3 结  论')
    # 云端 QA：Normal 缩进清零使全文缩短一页后，图 7/图 8 图带（显示第 12
    # 页底）的锚位若留在 S36 会滞留第 11 页触发屏障（右栏 12–15 cm 悬空，
    # 两次实测）。流位后移至结论首段之后，锚段随正文自然溢入第 12 页、与
    # 显示页同页（v9 原则）；图 7/图 8 首次引用均在前页，先文后图不变。
    emit_stream(doc, S4[:1], tables)
    add_floating_figure_band(doc, FIGS[7], FIGS[8])
    emit_stream(doc, S4[1:], tables)

    # 数据与程序可用性（贡献三的开源承诺；TXXB 无固定栏目，置结论后）
    p = doc.add_paragraph()
    pset(p, before=8, after=2, snap=False, line=260)
    emit_runs(p, '**数据与程序可用性：**', size=10.5)
    p = doc.add_paragraph()
    pset(p, snap=False, line=240)
    emit_runs(p, DATA_AVAIL, size=9)

    # 补充材料说明（附录 A–I 作为附件随投稿系统提交）
    p = doc.add_paragraph()
    pset(p, before=6, after=2, snap=False, line=260)
    emit_runs(p, '**补充材料说明：**', size=10.5)
    p = doc.add_paragraph()
    pset(p, snap=False, line=240)
    emit_runs(p, SUPP_NOTE, size=9)

    # 参考文献（TXXB 已刊体例：题头「参考文献 (References)」左顶格黑体；
    # 条目 6 号（8 磅），顺序编码制、悬挂缩进）
    p = doc.add_paragraph()
    pset(p, before=8, after=4, snap=False, line=260)
    emit_runs(p, '参考文献 (References)', size=10.5, ea=HEI, ascii_=HEI)
    missing_report = []
    for key, num in CITE_ORDER.items():
        if key not in entries:
            raise KeyError(f'引用键 {key} 不在 references.bib 中')
        etype, fields = entries[key]
        text, missing = gbt_format(key, etype, fields)
        if missing:
            missing_report.append((num, key, missing))
        p = doc.add_paragraph()
        pset(p, snap=False, line=220, ind_left=340, hanging=340)
        r = p.add_run(f'[{num}]\u2002')
        set_font(r, TNR, SONG, 8)
        r = p.add_run(text)
        set_font(r, TNR, SONG, 8)

    add_sect_break(doc, sect_2col)           # 结束双栏正文

    # ============ 末节：单栏（排版占位说明，投稿前删除） ============
    p = doc.add_paragraph()
    pset(p, before=10, after=2, snap=False, line=260)
    emit_runs(p, '**——以下为排版占位说明（投稿前请删除本块）——**', size=10.5)
    notes = [
        '1. 本文档为' + ('盲审稿（已删除中英文作者姓名、单位及作者简介，'
        '投稿指南 2(1)）' if blind else '原稿（保留中英文作者姓名、单位及'
        '作者简介，投稿指南 2(2)）') + '；两版本正文内容逐字一致。',
        '2. 全部【】括注为待作者补齐的占位符：作者姓名/单位/邮箱、基金项目名称及'
        '编号、第一作者与通信作者简介（姓名，出生年，性别，职称，学位，研究方向'
        '及邮箱，指南 1(3)）——以上均位于首页页脚与头区。',
        '3. 中图分类号现按模板默认「TP 391」填写，可视编辑部意见细化为 '
        'TP 391.41（图像识别与三维视觉）；DOI 与文章编号为模板占位样式，'
        '由编辑部赋号。',
        '4. 公式 (1)–(11) 与正文行内符号均为 Unicode 文本占位（已按“一式一号”'
        '拆分编号），投稿前需在 Word 中用公式编辑器/MathType 重排；变量斜体、'
        '矩阵/矢量黑斜体、标准函数正体的区分在重排时落实。',
        '5. 附录 A–I 及附图三张拟作为补充材料随投稿系统提交（见正文“补充材料'
        '说明”）；“版权转让协议”与“作者贡献及利益冲突声明”两个 PDF 需另行'
        '下载模板签署上传（指南 3(2)），正文中不含作者贡献声明。',
        '6. 参考文献会议条目：会议录全称与出版地/出版者按通行著录填制，个别'
        '小型会议以 [S.l.]: [s.n.] 占位，投稿前请作者终核。',
        '7. 篇幅风险：投稿指南 1(1) 称每篇以 5 000～8 000 字为宜，本文'
        '（含图表）明显超出该体量；转刊未删任何内容，是否压缩交作者决策'
        '（可参照 oep_submission_20260827/REVIEW 报告的附录外置与大表'
        '瘦身路线）。费用参考：稿件处理费约 100 元/篇 + 出版服务费及彩图'
        '制作费约 500 元/版（指南 5）。',
    ]
    for note in notes:
        p = doc.add_paragraph()
        pset(p, snap=False, line=240, after=2)
        emit_runs(p, note, size=9)

    # ============ 首页脚注区（TXXB 模板 first footer） ============
    # 模板自带样式：收稿/定稿日期双语行 + 基金项目双语行；按投稿指南 1(3)
    # 与模板正文说明补第一作者及通信作者简介行（6 号宋体）。盲审稿删除
    # 作者简介与通信作者行，保留收稿日期与基金项目占位。
    sec0 = doc.sections[0]
    try:
        fp = sec0.first_page_footer
        for par in list(fp.paragraphs):
            par._p.getparent().remove(par._p)
        lines = [
            '收稿日期：0000-00-00；定稿日期：0000-00-00（由编辑部/投稿系统填写）',
            '基金项目：【基金项目名称及编号，如：国家自然科学基金项目(00000000)】',
        ]
        if not blind:
            lines += [
                '第一作者：【姓名】（【出生年】），【性别】，【职称】，【学位】。'
                '主要研究方向为【…】。E-mail：【email@example.edu.cn】',
                '通信作者：【姓名】（【出生年】），【性别】，【职称】，【学位】。'
                '主要研究方向为【…】。E-mail：【email@example.edu.cn】',
            ]
        else:
            lines += [' ', ' ']   # 等高占位，保持首页版面与原稿一致
        first = True
        for text in lines:
            par = fp.add_paragraph()
            if first:
                # 脚注区首行加上边线，与正文区分（模板体例）
                pPr = par._p.get_or_add_pPr()
                pbdr = OxmlElement('w:pBdr')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '6')
                top.set(qn('w:space'), '1')
                top.set(qn('w:color'), 'auto')
                pbdr.append(top)
                pPr.append(pbdr)
                first = False
            pset(par, snap=False, line=200)
            emit_runs(par, text, size=7.5, ea=SONG)
    except Exception as e:      # 页脚结构与预期不符时不阻塞主流程
        print('[warn] 首页脚注区重写跳过：', e)

    out = OUT_BLIND if blind else OUT_FULL
    doc.save(out)

    # ============ 自检 ============
    chk = Document(out)
    n_par = len(chk.paragraphs)
    n_tbl = len(chk.tables)
    n_img = len(chk.inline_shapes)
    n_sect = len(chk.sections)
    ver = '盲审稿' if blind else '原稿'
    print(f'[ok] 已生成（{ver}）{out}')
    print(f'     段落 {n_par} | 顶级表格 {n_tbl}（内嵌三线表 7 + 单图浮动容器 4 '
          f'+ 双图图带 2）| 图片 {n_img} | 分节 {n_sect} | '
          f'编号公式 {EQN_COUNT[0]} | 参考文献 {len(CITE_ORDER)}')
    assert n_img == 8, f'图片数应为 8，实际 {n_img}'
    assert n_tbl == 13, \
        f'顶级表格数应为 13（内嵌三线表 7 + 单图容器 4 + 图带 2），' \
        f'实际 {n_tbl}'
    n_nested = sum(len(c.tables) for t in chk.tables
                   for row in t.rows for c in row.cells)
    assert n_nested == 0, f'v11 无容器内嵌套三线表，实际 {n_nested}'
    assert n_sect == 5, f'分节数应为 5（头部/表区/末节分节），实际 {n_sect}'
    assert len(CITE_ORDER) == 55, f'引用数应为 55，实际 {len(CITE_ORDER)}'
    assert EQN_COUNT[0] == 11, f'编号公式应为 11（一式一号），实际 {EQN_COUNT[0]}'
    # 盲审自检：正文与页脚不得残留作者/单位占位
    if blind:
        full_text = '\n'.join(p.text for p in chk.paragraphs)
        for s in chk.sections:
            full_text += '\n'.join(p.text for p in s.first_page_footer.paragraphs)
        for token in ('作者一姓名', 'SURNAME', '第一单位', 'Department, University',
                      '第一作者：', '通信作者：'):
            assert token not in full_text, f'盲审稿残留身份信息占位：{token}'
        print('     盲审自检通过：无作者/单位/简介占位残留')
    if missing_report:
        print('     [文献待补字段]')
        for num, key, missing in missing_report:
            print(f'       [{num}] {key}: {"、".join(sorted(set(missing)))}')
    size_mb = os.path.getsize(out) / 1e6
    print(f'     文件大小 {size_mb:.2f} MB')
    return missing_report


def _reset_counters():
    """两个版本连续构建前重置模块级计数器（引用编号表与公式计数）。"""
    CITE_ORDER.clear()
    EQN_COUNT[0] = 0


if __name__ == '__main__':
    if '--blind-only' in sys.argv:
        _reset_counters()
        build(blind=True)
    elif '--full-only' in sys.argv:
        _reset_counters()
        build(blind=False)
    else:
        _reset_counters()
        build(blind=False)
        _reset_counters()
        build(blind=True)
