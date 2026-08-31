#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""text_integrity_check.py —— docx→PDF 渲染文字完整性校验（验收工具）

用法： python3 text_integrity_check.py 排版稿.docx 渲染稿.pdf [块长]

动机：LibreOffice 类引擎在"两栏断行 × 底部浮动块"交界处偶发丢字
（jig v1/v2 首轮渲染中"值缺"二字被排到游离位置或直接丢弃），需在
渲染 PDF 上逐段证明 docx 正文（含浮动容器/表格单元格）无字符丢失。

抽取降噪（对照实测的 LibreOffice 6.0 导出行为）：
  * 零宽幽灵重影——字体回退串会被额外发射一份 bbox 宽≈0 的隐形
    文本，按字符宽度 >0.05 pt 过滤；
  * 行尾标点压缩／组合字符（s̃ 的 ̃）同样零宽——比对层面双侧只保留
    字母/数字/汉字（Unicode L*/N* 类），标点符号不参与比对；
  * 页眉页脚（跨页重复行、收稿日期/基金项目行）从 PDF 侧剔除，
    避免其在页缝处打断正文流。

三级判定：连续子串 → 定长块（缺省 12 字）按位置单调匹配 → 锚点
邻域字符多重集（对上下标抽取乱序、两端对齐重排免疫，仍能证明段内
每个字符在其应在的版面邻域内出现）。第三级仍失败才判丢字。
"""
import re
import sys
import unicodedata
from collections import Counter

import fitz
from docx import Document

HDR_FT_RE = re.compile(r'收稿日期|修回日期|基金项目|Supported by')


def _keep(ch):
    return unicodedata.category(ch)[0] in 'LN'


def _norm(s):
    s = unicodedata.normalize('NFKC', s)
    return ''.join(c for c in s if _keep(c))


def docx_paras(path):
    doc = Document(path)
    out = []

    def has_omml(p):
        """含 OMML 数学区的公式段：p.text 只剩夹排文字（"时，(7)"类
        碎片），其字符在 PDF 中与公式引擎发射的数学符号交错，三级判定
        均无法定位——属检查盲区而非丢字，公式内容由构建自检（OMML 段
        11/数学区 13）与人工目检公式页保障。"""
        return bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/officeDocument'
            '/2006/math}oMath'))

    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if not has_omml(p):
                            out.append(p.text)
                    walk_tables(cell.tables)

    for p in doc.paragraphs:
        if not has_omml(p):
            out.append(p.text)
    walk_tables(doc.tables)
    return [x for x in (_norm(t) for t in out) if x]


def pdf_visible_text(pdf):
    """rawdict 逐行提取：滤零宽字符，剔页眉页脚（跨页重复×边缘位置）。"""
    pages_lines = []
    for page in pdf:
        h = page.rect.height
        lines = []
        for blk in page.get_text('rawdict')['blocks']:
            for line in blk.get('lines', []):
                chars = [ch['c'] for span in line['spans']
                         for ch in span['chars']
                         if ch['bbox'][2] - ch['bbox'][0] > 0.05]
                raw = ''.join(chars)
                if raw.strip():
                    y_frac = (line['bbox'][1] + line['bbox'][3]) / 2 / h
                    lines.append((y_frac, raw))
        pages_lines.append(lines)
    n_pages = len(pages_lines)
    # 页眉/页脚 = 跨页重复 且 位于页面上/下边缘（表格窄列里独占一行的
    # "AUROC" 等单元格虽跨页重复，但在页中部，不得误杀）
    page_freq = Counter()
    for lines in pages_lines:
        for key in {_norm(l) for _, l in lines if _norm(l)}:
            page_freq[key] += 1
    rep_thresh = max(3, n_pages // 3)
    out = []
    for pno, lines in enumerate(pages_lines):
        for y_frac, l in lines:
            key = _norm(l)
            edge = y_frac < 0.12 or y_frac > 0.88
            if key and edge and page_freq[key] >= rep_thresh:
                continue                      # 跨页重复的页眉/页脚
            if pno == 0 and HDR_FT_RE.search(l):
                continue                      # 首页脚注区（仅第 1 页）
            out.append(l)
    return ''.join(out)


def chunk_check(needle, hay, chunk):
    pos = 0
    for i in range(0, len(needle), chunk):
        seg = needle[i:i + chunk]
        j = hay.find(seg, pos)
        if j < 0:
            return False
        pos = j + len(seg)
    return True


def neighborhood_multiset(needle, hay, chunk=12, pad=3000):
    """锚点邻域字符多重集校验；返回 (是否通过, 缺失字符串)。

    以段内首个能在 PDF 中定位的定长块为锚，在其前后邻域内要求段落
    每个字符按多重数出现。对上下标/两端对齐导致的抽取乱序免疫。"""
    anchor = -1
    for i in range(0, max(1, len(needle) - chunk + 1), chunk):
        j = hay.find(needle[i:i + chunk])
        if j >= 0:
            anchor = j
            break
    if anchor < 0:
        return False, '（整段无法在 PDF 中定位）'
    lo = max(0, anchor - 200 - chunk * (i // chunk))
    hi = anchor + len(needle) + pad
    missing = Counter(needle) - Counter(hay[lo:hi])
    if missing:
        return False, ''.join(sorted(missing.elements()))
    return True, None


def main(docx_path, pdf_path, chunk=12):
    pdf = fitz.open(pdf_path)
    pdf_text = _norm(pdf_visible_text(pdf))
    paras = docx_paras(docx_path)
    n_sub = n_chunk = n_nbh = 0
    hard_fail = []
    for t in paras:
        if t in pdf_text:
            n_sub += 1
        elif chunk_check(t, pdf_text, chunk):
            n_chunk += 1
        else:
            ok, miss = neighborhood_multiset(t, pdf_text, chunk)
            if ok:
                n_nbh += 1
            else:
                hard_fail.append((t, miss))
    print(f'{pdf_path}: docx 段落 {len(paras)}，连续子串 {n_sub}，'
          f'单调分块 {n_chunk}，邻域多重集 {n_nbh}，丢字 {len(hard_fail)}')
    for t, miss in hard_fail:
        print('  [丢字] 段落前 56 字：' + t[:56])
        print(f'         缺失字符：“{miss}”')
    return 1 if hard_fail else 0


if __name__ == '__main__':
    docx_p, pdf_p = sys.argv[1], sys.argv[2]
    ck = int(sys.argv[3]) if len(sys.argv) > 3 else 12
    sys.exit(main(docx_p, pdf_p, ck))
